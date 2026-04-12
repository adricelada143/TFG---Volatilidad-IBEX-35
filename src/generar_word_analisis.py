"""
Genera el archivo Word de la memoria de Análisis del Dato
con todas las figuras y tablas estilizadas como imágenes.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

# ── Rutas ──────────────────────────────────────────────────────────
BASE = "/Users/adriancelada/Library/Mobile Documents/com~apple~CloudDocs/UFV/UNIVERSIDAD FRANCISCO DE VITORIA/4º/TFG"
FIGS = os.path.join(BASE, "proyecto/outputs/figuras")
TABS = os.path.join(BASE, "proyecto/outputs/tablas_img_analisis")
OUT  = os.path.join(BASE, "Análisis del Dato/Memoria_Analisis_del_Dato.docx")

os.makedirs(TABS, exist_ok=True)
os.makedirs(os.path.dirname(OUT), exist_ok=True)

doc = Document()

# ── Estilos ────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0, 51, 102)
    if level == 1:
        hs.font.size = Pt(16)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(13)
        hs.font.bold = True
    else:
        hs.font.size = Pt(11)
        hs.font.bold = True
        hs.font.italic = True


# ── Funciones auxiliares ───────────────────────────────────────────
def add_paragraph(text, bold=False, italic=False, style_name='Normal'):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_figure(filename, caption, width=Inches(5.8)):
    path = os.path.join(FIGS, filename)
    if not os.path.exists(path):
        add_paragraph(f"[Figura no encontrada: {filename}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(80, 80, 80)


def render_table_image(headers, rows, filename, title=None, col_widths=None,
                       highlight_col=None, highlight_vals=None):
    """
    Renderiza una tabla como imagen PNG estilizada.
    - headers: lista de strings
    - rows: lista de listas de strings
    - highlight_col: índice de columna para colorear celdas condicionalmente
    - highlight_vals: dict {valor: color_hex} para resaltar celdas
    """
    n_rows = len(rows)
    n_cols = len(headers)

    # Calcular anchos automáticos basados en contenido
    if col_widths is None:
        max_lens = []
        for c in range(n_cols):
            max_len = len(str(headers[c]))
            for r in rows:
                max_len = max(max_len, len(str(r[c])))
            max_lens.append(max_len)
        total = sum(max_lens)
        col_widths = [m / total for m in max_lens]

    fig_width = max(6, min(12, n_cols * 1.8))
    row_height = 0.38
    fig_height = (n_rows + 1) * row_height + (0.5 if title else 0.15) + 0.15

    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, n_rows + 1)
    ax.axis('off')

    # Colores
    HEADER_BG = '#003366'
    HEADER_FG = '#FFFFFF'
    ROW_EVEN  = '#F7F9FC'
    ROW_ODD   = '#FFFFFF'
    BORDER    = '#D0D5DD'
    TEXT_COLOR = '#1A1A2E'

    # Título
    if title:
        fig.suptitle(title, fontsize=10, fontweight='bold', color='#003366',
                     y=1 - 0.08 / fig_height, ha='center')

    # Dibujar celdas
    for row_idx in range(n_rows + 1):  # +1 para header
        y = n_rows - row_idx  # invertir para que header arriba

        for col_idx in range(n_cols):
            x_start = sum(col_widths[:col_idx])
            w = col_widths[col_idx]

            # Background
            if row_idx == 0:
                bg = HEADER_BG
                fg = HEADER_FG
                fw = 'bold'
                fs = 8.5
            else:
                bg = ROW_EVEN if row_idx % 2 == 0 else ROW_ODD
                fg = TEXT_COLOR
                fw = 'normal'
                fs = 8

                # Highlight condicional
                if highlight_col is not None and highlight_vals and col_idx == highlight_col:
                    val = str(rows[row_idx - 1][col_idx]).strip()
                    if val in highlight_vals:
                        bg = highlight_vals[val]

            rect = plt.Rectangle((x_start, y), w, 1,
                                 facecolor=bg, edgecolor=BORDER, linewidth=0.5)
            ax.add_patch(rect)

            # Texto
            text = headers[col_idx] if row_idx == 0 else str(rows[row_idx - 1][col_idx])
            ax.text(x_start + w / 2, y + 0.5, text,
                    ha='center', va='center', fontsize=fs,
                    color=fg, fontweight=fw, fontfamily='sans-serif')

    plt.tight_layout(pad=0.2)
    path = os.path.join(TABS, filename)
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(fig)
    return path


def add_table_img(headers, rows, filename, caption=None, title=None,
                  width=Inches(5.5), highlight_col=None, highlight_vals=None,
                  col_widths=None):
    """Renderiza tabla como imagen y la inserta en el doc."""
    path = render_table_image(headers, rows, filename, title=title,
                              highlight_col=highlight_col,
                              highlight_vals=highlight_vals,
                              col_widths=col_widths)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(80, 80, 80)


# ════════════════════════════════════════════════════════════════════
#  CONTENIDO
# ════════════════════════════════════════════════════════════════════

# ── TÍTULO ─────────────────────────────────────────────────────────
title = doc.add_heading('Análisis del Dato', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 1. MARCO TEÓRICO
# ════════════════════════════════════════════════════════════════════
doc.add_heading('1. Marco Teórico', level=1)

add_paragraph(
    'La predicción de la volatilidad en los mercados financieros constituye uno de los problemas '
    'centrales de la econometría financiera moderna. Desde los trabajos seminales de Mandelbrot (1963), '
    'quien documentó por primera vez el fenómeno de volatility clustering —la tendencia de los períodos '
    'de alta volatilidad a agruparse en el tiempo—, la comunidad académica ha desarrollado un amplio '
    'arsenal de modelos que intentan capturar la dinámica temporal de la varianza condicional de los '
    'rendimientos financieros.'
)

add_paragraph(
    'El marco teórico de este trabajo se articula en torno a cuatro pilares fundamentales que reflejan '
    'la evolución histórica del campo y justifican las decisiones metodológicas adoptadas.'
)

doc.add_heading('1.1. Modelos de varianza condicional: de ARCH a GARCH', level=2)
add_paragraph(
    'Engle (1982) propuso el modelo ARCH (Autoregressive Conditional Heteroskedasticity), que formalizó '
    'la idea de que la varianza de los rendimientos financieros no es constante sino que depende de los '
    'shocks pasados. Este modelo captura la observación empírica de que retornos grandes (en valor absoluto) '
    'tienden a ser seguidos por retornos grandes, independientemente de su signo. Formalmente, el modelo '
    'ARCH(q) especifica la varianza condicional como:'
)
add_paragraph('    σ²ₜ = ω + α₁·ε²ₜ₋₁ + α₂·ε²ₜ₋₂ + … + αq·ε²ₜ₋q')
add_paragraph(
    'Bollerslev (1986) generalizó este enfoque al modelo GARCH (Generalized ARCH), que incorpora '
    'retardos de la propia varianza condicional. El modelo GARCH(1,1), que se emplea en este trabajo, '
    'especifica:'
)
add_paragraph('    σ²ₜ = ω + α·ε²ₜ₋₁ + β·σ²ₜ₋₁')
add_paragraph(
    'donde α captura el impacto de los shocks recientes y β la persistencia de la volatilidad. '
    'La condición de estacionariedad requiere α + β < 1, y valores cercanos a 1 indican alta '
    'persistencia —un fenómeno ubicuo en los mercados financieros. El modelo GARCH(1,1) con '
    'innovaciones t-Student se ha convertido en el estándar de la industria para la modelización de '
    'volatilidad, por su capacidad para capturar simultáneamente el clustering de volatilidad y las '
    'colas pesadas de la distribución de rendimientos.'
)

doc.add_heading('1.2. El modelo HAR para volatilidad realizada', level=2)
add_paragraph(
    'Corsi (2009) propuso el modelo HAR (Heterogeneous Autoregressive), que adopta un enfoque '
    'radicalmente diferente al GARCH. En lugar de modelizar la varianza condicional a través de un '
    'proceso autorregresivo latente, el HAR predice directamente la volatilidad realizada usando '
    'retardos a tres escalas temporales:'
)
add_paragraph('    Vₜ = β₀ + β₁·Vₜ₋₁ + β₂·Vₜ₋₅ + β₃·Vₜ₋₂₁ + εₜ')
add_paragraph(
    'donde Vₜ₋₁, Vₜ₋₅ y Vₜ₋₂₁ representan la volatilidad promedio en los últimos 1, 5 y 21 días '
    'de negociación (escalas diaria, semanal y mensual, respectivamente). La motivación teórica del HAR '
    'reside en la hipótesis del mercado heterogéneo: distintos tipos de agentes (traders intradía, '
    'gestores semanales, inversores institucionales mensuales) operan en escalas temporales distintas, '
    'y la volatilidad agregada refleja la superposición de estas escalas.'
)
add_paragraph(
    'El modelo HAR, pese a su simplicidad (es una regresión lineal), ha demostrado un rendimiento '
    'predictivo competitivo con modelos mucho más complejos, lo que lo convierte en un benchmark natural '
    'para la predicción de volatilidad.'
)

doc.add_heading('1.3. Machine Learning: XGBoost y relaciones no lineales', level=2)
add_paragraph(
    'Chen y Guestrin (2016) presentaron XGBoost (eXtreme Gradient Boosting), un algoritmo de aprendizaje '
    'por ensamble basado en gradient boosting de árboles de decisión. Su principal ventaja frente a los '
    'modelos lineales es la capacidad de capturar relaciones no lineales e interacciones entre variables '
    'sin necesidad de especificarlas explícitamente. En el contexto de predicción de volatilidad, '
    'XGBoost puede detectar:'
)
items = [
    'Efectos umbral: el VIX puede tener un impacto despreciable por debajo de 20 pero muy fuerte por encima de 30.',
    'Interacciones: la combinación de tipos de interés altos con PIB negativo puede amplificar la volatilidad más que la suma de sus efectos individuales.',
    'No linealidades: la relación entre el spread soberano y la volatilidad puede ser convexa (crece aceleradamente).',
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)
add_paragraph(
    'Lundberg y Lee (2017) propusieron los valores SHAP (SHapley Additive exPlanations), basados en la '
    'teoría de juegos cooperativos de Shapley, que permiten descomponer la predicción de cada observación '
    'en las contribuciones individuales de cada feature. Los valores SHAP proporcionan interpretabilidad '
    'local y global, resolviendo el problema clásico de la "caja negra" del machine learning.'
)

doc.add_heading('1.4. Evaluación de modelos de volatilidad', level=2)
add_paragraph(
    'La evaluación de modelos de predicción de volatilidad requiere métricas específicas. Patton (2011) '
    'demostró que las funciones de pérdida habituales (MSE, MAE) pueden producir rankings de modelos '
    'inconsistentes cuando la volatilidad verdadera no es observable —como ocurre en la práctica, donde '
    'solo se dispone de un proxy (la volatilidad realizada). Patton identificó las funciones de pérdida '
    'que son robustas a este problema, entre las que destaca QLIKE:'
)
add_paragraph('    QLIKE = mean(yₜ / ŷₜ − ln(yₜ / ŷₜ) − 1)')
add_paragraph(
    'donde yₜ es la volatilidad observada y ŷₜ la predicha. QLIKE penaliza asimétricamente: '
    'subestimar la volatilidad se castiga más que sobreestimarla, lo que tiene sentido económico '
    'en gestión de riesgos.'
)
add_paragraph(
    'Para la comparación formal entre modelos, Diebold y Mariano (1995) propusieron un test estadístico '
    '(DM test) que evalúa si la diferencia en la capacidad predictiva de dos modelos es estadísticamente '
    'significativa, teniendo en cuenta la autocorrelación de los errores de predicción.'
)

doc.add_heading('1.5. Contexto empírico: el IBEX 35 y las variables macroeconómicas', level=2)
add_paragraph(
    'El IBEX 35 presenta características específicas que lo distinguen de otros índices bursátiles '
    'europeos. Su elevada concentración sectorial en banca (Santander, BBVA, CaixaBank) y utilities '
    '(Iberdrola, Endesa) lo hace especialmente sensible a las decisiones del BCE sobre tipos de interés. '
    'Además, la crisis de deuda soberana de 2010-2012 dejó una huella particular en la relación entre '
    'la prima de riesgo española y la volatilidad bursátil.'
)
add_paragraph(
    'Este trabajo se articula en torno a dos preguntas de investigación complementarias:'
)
p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Pregunta 1 (Predicción): ')
run_b.bold = True
p.add_run(
    '¿Pueden las variables macroeconómicas mejorar la predicción de la volatilidad del IBEX 35 '
    'más allá de lo que ya captura la propia historia de la volatilidad?'
)
p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Pregunta 2 (Event Study): ')
run_b.bold = True
p.add_run(
    '¿Las fechas de publicación de los indicadores macroeconómicos generan shocks significativos '
    'en la volatilidad intradía del IBEX 35?'
)
add_paragraph(
    'La distinción entre ambas preguntas es crucial: la primera evalúa si los niveles de las variables '
    'macro predicen la volatilidad futura; la segunda evalúa si el acto de publicación (independientemente '
    'del nivel) genera una reacción inmediata en el mercado.'
)

# ════════════════════════════════════════════════════════════════════
# 2. CAPÍTULO 1: MODELOS PREDICTIVOS DE VOLATILIDAD
# ════════════════════════════════════════════════════════════════════
doc.add_heading('2. Capítulo 1: Modelos Predictivos de Volatilidad', level=1)

# 2.1
doc.add_heading('2.1. Variable objetivo', level=2)
add_paragraph(
    'La variable objetivo de los modelos predictivos es la volatilidad histórica anualizada a 21 días '
    '(vol_hist_21d), definida como:'
)
add_paragraph('    σₜ,₂₁ = √252 · std(rₜ₋₂₀, rₜ₋₁₉, …, rₜ)')
add_paragraph(
    'donde rₜ representa el log-retorno diario y 252 es el número estándar de días de negociación al año. '
    'La ventana de 21 días se elige por convención del mercado (aproximadamente un mes bursátil) y porque '
    'proporciona un balance adecuado entre suavizado de ruido y capacidad de reacción ante cambios de régimen.'
)
add_paragraph(
    'El dataset utilizado para la modelización contiene 3.455 observaciones diarias del IBEX 35 '
    'agregado (mediana de las 35 empresas), cubriendo el período de enero de 2012 a octubre de 2025. '
    'Se adopta un split temporal 80/20: las primeras 2.764 observaciones (hasta aproximadamente '
    'mediados de 2023) conforman el conjunto de entrenamiento, y las 691 restantes el conjunto de test. '
    'El split es estrictamente temporal (no aleatorio) para evitar data leakage y reflejar el escenario '
    'real de predicción fuera de muestra.'
)
add_table_img(
    ['Concepto', 'Valor'],
    [
        ['Observaciones totales', '3.455'],
        ['Período', 'Enero 2012 – Octubre 2025'],
        ['Train (80%)', '2.764 observaciones'],
        ['Test (20%)', '691 observaciones'],
        ['Tipo de split', 'Temporal (cronológico)'],
        ['Variable objetivo', 'vol_hist_21d (mediana IBEX 35)'],
    ],
    'tab_2_1_split.png',
    caption='Tabla 1. Configuración del dataset de modelización y split train/test.',
    col_widths=[0.45, 0.45]
)
add_figure('15_split_temporal.png',
           'Figura 15. Visualización del split temporal train/test. La línea vertical marca la frontera entre '
           'entrenamiento (azul) y test (naranja).')

# 2.2
doc.add_heading('2.2. Features: variables HAR y macroeconómicas', level=2)
add_paragraph(
    'El conjunto de features se organiza en dos bloques siguiendo la lógica del modelo HAR (Corsi, 2009) '
    'aumentado con variables macroeconómicas:'
)
add_paragraph('Bloque 1 — Variables HAR (endógenas):', bold=True)
items_har = [
    'vol_lag1: volatilidad del día anterior (escala diaria).',
    'vol_lag5: volatilidad promedio de los últimos 5 días (escala semanal).',
    'vol_lag21: volatilidad promedio de los últimos 21 días (escala mensual).',
]
for item in items_har:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)
add_paragraph(
    'Estas tres variables capturan la memoria larga de la volatilidad documentada en el análisis '
    'exploratorio (ACF significativa en todos los lags hasta 60). Constituyen el núcleo predictivo '
    'de todos los modelos.'
)

add_paragraph('Bloque 2 — Variables macroeconómicas (exógenas):', bold=True)
add_paragraph(
    'Se incluyen 12 variables macroeconómicas organizadas en tres sub-bloques temáticos que reflejan '
    'distintos canales de transmisión hacia la volatilidad bursátil:'
)
add_table_img(
    ['Sub-bloque', 'Variables', 'Canal'],
    [
        ['Riesgo global', 'vix, vstoxx, brent, gas_ttf', 'Contagio entre mercados'],
        ['Condiciones monetarias', 'bono_es_10y, bono_de_10y, spread_es_de, euribor_3m, eur_usd, tipo_dfr', 'Política monetaria'],
        ['Actividad económica', 'ipc_yoy, tasa_paro', 'Ciclo real'],
    ],
    'tab_2_2_features.png',
    caption='Tabla 2. Variables macroeconómicas agrupadas por canal de transmisión.',
    col_widths=[0.22, 0.48, 0.25]
)
add_paragraph(
    'En total, cada modelo dispone de 15 features potenciales (3 HAR + 12 macro). Los modelos "solo HAR" '
    'utilizan únicamente las 3 primeras; los modelos "HAR + Macro" emplean las 15.'
)

# 2.3
doc.add_heading('2.3. Métricas de evaluación', level=2)
add_paragraph(
    'La evaluación de los modelos se realiza con cuatro métricas complementarias, cada una capturando '
    'una dimensión distinta de la calidad predictiva. Las tres primeras son métricas estándar de regresión; '
    'la cuarta es específica para modelos de volatilidad.'
)

add_paragraph('RMSE (Root Mean Squared Error):', bold=True)
add_paragraph('    RMSE = √(1/n · Σ(yₜ − ŷₜ)²)')
add_paragraph(
    'Penaliza cuadráticamente los errores grandes. Es la métrica principal de comparación porque está '
    'en las mismas unidades que la variable objetivo (volatilidad anualizada) y es sensible a errores '
    'en episodios de alta volatilidad, que son los más relevantes para la gestión de riesgos.'
)

add_paragraph('MAE (Mean Absolute Error):', bold=True)
add_paragraph('    MAE = 1/n · Σ|yₜ − ŷₜ|')
add_paragraph(
    'Proporciona una medida del error medio sin penalización cuadrática. Es más robusta a outliers que '
    'el RMSE y refleja el error "típico" del modelo en condiciones normales de mercado.'
)

add_paragraph('R² (Coeficiente de determinación):', bold=True)
add_paragraph('    R² = 1 − SS_res / SS_tot = 1 − Σ(yₜ − ŷₜ)² / Σ(yₜ − ȳ)²')
add_paragraph(
    'Mide la proporción de varianza explicada. Un R² de 0,97 indica que el modelo explica el 97% de la '
    'variabilidad de la volatilidad. Puede ser negativo si el modelo es peor que la media histórica como '
    'predictor, lo que ocurrirá con el GARCH por razones que se detallan en la sección 2.6.'
)

add_paragraph('QLIKE (Quasi-Likelihood Loss — Patton, 2011):', bold=True)
add_paragraph('    QLIKE = 1/n · Σ(yₜ/ŷₜ − ln(yₜ/ŷₜ) − 1)')
add_paragraph(
    'Función de pérdida robusta específica para volatilidad, derivada por Patton (2011) como una de las '
    'pocas funciones que producen rankings de modelos consistentes cuando la volatilidad verdadera se '
    'aproxima por un proxy ruidoso. Valores más negativos indican mejor ajuste. QLIKE penaliza la '
    'subestimación de volatilidad (yₜ/ŷₜ >> 1) más que la sobreestimación, alineándose con el principio '
    'de prudencia en gestión de riesgos.'
)

# 2.4
doc.add_heading('2.4. Modelo A0: Regresión Simple (Baseline)', level=2)
add_paragraph(
    'El modelo más simple posible sirve como línea base (baseline) contra la que evaluar los modelos '
    'más complejos. Se trata de una regresión lineal univariante que utiliza únicamente la volatilidad '
    'del día anterior para predecir la volatilidad del día siguiente:'
)
add_paragraph('    V̂ₜ = β₀ + β₁ · Vₜ₋₁')
add_paragraph(
    'Este modelo captura la idea fundamental de persistencia de la volatilidad: el mejor predictor de la '
    'volatilidad de mañana es la volatilidad de hoy. Pese a su extrema simplicidad (un único regresor), '
    'este modelo establece un listón sorprendentemente alto gracias a la fuerte autocorrelación de la '
    'serie de volatilidad documentada en el EDA.'
)
add_paragraph('Resultados en el conjunto de test:', bold=True)
add_table_img(
    ['Métrica', 'Valor'],
    [
        ['RMSE', '0,0083'],
        ['MAE', '0,0048'],
        ['R²', '0,9686'],
        ['QLIKE', '−1,8419'],
    ],
    'tab_2_4_simple.png',
    caption='Tabla 3. Resultados del Modelo A0 (Regresión Simple) en test.',
    col_widths=[0.4, 0.4]
)
add_paragraph(
    'Con un R² de 0,969, la regresión simple explica casi el 97% de la varianza de la volatilidad '
    'utilizando un solo lag. Este resultado confirma la dominancia de la persistencia de volatilidad '
    'como mecanismo predictivo y establece un umbral exigente para los modelos más complejos.'
)
add_figure('16b_simple_resultados.png',
           'Figura 16. Modelo A0 (Regresión Simple): predicción vs. valores reales en el conjunto de test.')
add_figure('16d_simple_diagnostico.png',
           'Figura 17. Modelo A0 (Regresión Simple): diagnóstico de residuos — histograma, QQ-plot, ACF y error temporal.')

# 2.5
doc.add_heading('2.5. Modelo A: OLS HAR + Macro', level=2)
add_paragraph(
    'El Modelo A es una regresión lineal por mínimos cuadrados ordinarios (OLS) que incorpora las 15 '
    'features disponibles: 3 variables HAR y 12 variables macroeconómicas. Representa la versión '
    'completa del modelo HAR de Corsi (2009) aumentado con información macroeconómica:'
)
add_paragraph(
    '    V̂ₜ = β₀ + β₁·vol_lag1 + β₂·vol_lag5 + β₃·vol_lag21 + β₄·vix + β₅·vstoxx + '
    'β₆·brent + β₇·gas_ttf + β₈·bono_es_10y + β₉·bono_de_10y + β₁₀·spread_es_de + '
    'β₁₁·euribor_3m + β₁₂·eur_usd + β₁₃·tipo_dfr + β₁₄·ipc_yoy + β₁₅·tasa_paro'
)
add_paragraph('Coeficientes estimados:', bold=True)
add_table_img(
    ['Variable', 'Coeficiente', 'Variable', 'Coeficiente'],
    [
        ['Intercepto', '0,0229', 'brent', '0,0000'],
        ['vol_lag1', '0,8982', 'gas_ttf', '−0,0000'],
        ['vol_lag5', '0,0193', 'bono_es_10y', '0,0016'],
        ['vol_lag21', '0,0464', 'bono_de_10y', '−0,0006'],
        ['vix', '0,0002', 'spread_es_de', '0,0000'],
        ['vstoxx', '−0,0001', 'euribor_3m', '−0,0010'],
        ['eur_usd', '−0,0121', 'tipo_dfr', '0,0003'],
        ['ipc_yoy', '−0,0002', 'tasa_paro', '0,0000'],
    ],
    'tab_2_5_coefs_ols.png',
    caption='Tabla 4. Coeficientes estimados del Modelo A (OLS HAR + Macro).',
    col_widths=[0.22, 0.22, 0.22, 0.22]
)
add_paragraph(
    'Se observa que vol_lag1 domina con un coeficiente de 0,898 (consistente con la alta persistencia), '
    'mientras que los coeficientes de las variables macro son de magnitud muy pequeña, sugiriendo un '
    'efecto marginal limitado en el marco lineal.'
)
add_paragraph('Resultados en el conjunto de test:', bold=True)
add_table_img(
    ['Métrica', 'Valor'],
    [
        ['RMSE', '0,0084'],
        ['MAE', '0,0061'],
        ['R²', '0,9677'],
        ['QLIKE', '−1,8416'],
    ],
    'tab_2_5_resultados_ols.png',
    caption='Tabla 5. Resultados del Modelo A (OLS HAR + Macro) en test.',
    col_widths=[0.4, 0.4]
)
add_paragraph(
    'El OLS con 15 features obtiene un RMSE de 0,0084, ligeramente peor que la regresión simple (0,0083). '
    'Esto indica que, en el marco lineal, la adición de variables macroeconómicas no mejora la predicción '
    'sino que introduce ruido, un hallazgo consistente con el riesgo de sobreajuste cuando se añaden '
    'features de bajo poder predictivo marginal.'
)

add_paragraph('Diagnóstico de residuos:', bold=True)
add_table_img(
    ['Estadístico', 'Valor', 'Interpretación'],
    [
        ['Jarque-Bera', '4.895,6', 'Fuerte rechazo de normalidad'],
        ['Asimetría (skewness)', '−0,209', 'Ligera cola izquierda'],
        ['Curtosis (exceso)', '13,137', 'Colas muy pesadas'],
    ],
    'tab_2_5_diagnostico_ols.png',
    caption='Tabla 6. Diagnóstico de residuos del Modelo A (OLS HAR + Macro).',
    col_widths=[0.3, 0.2, 0.4]
)
add_paragraph(
    'Los residuos presentan colas pesadas (curtosis de 13,1), lo que indica que el modelo lineal no '
    'captura adecuadamente los episodios de volatilidad extrema. Sin embargo, las estimaciones puntuales '
    'de los coeficientes siguen siendo consistentes por el teorema de Gauss-Markov, aunque los intervalos '
    'de confianza deben interpretarse con cautela.'
)
add_figure('16_ols_resultados.png',
           'Figura 18. Modelo A (OLS HAR + Macro): predicción vs. valores reales en el conjunto de test.')
add_figure('16c_ols_diagnostico.png',
           'Figura 19. Modelo A (OLS HAR + Macro): diagnóstico de residuos — histograma, QQ-plot, ACF y error temporal.')

# 2.6
doc.add_heading('2.6. Modelo B: GARCH(1,1) con innovaciones t-Student', level=2)
add_paragraph(
    'El modelo GARCH(1,1) adopta un enfoque fundamentalmente distinto a los modelos de regresión. '
    'En lugar de predecir la volatilidad realizada directamente, modela la varianza condicional del '
    'proceso generador de retornos. La especificación completa es:'
)
add_paragraph('Ecuación de la media:', bold=True)
add_paragraph('    rₜ = μ + εₜ,    donde εₜ = σₜ · zₜ,    zₜ ~ t(ν)')
add_paragraph('Ecuación de la varianza:', bold=True)
add_paragraph('    σ²ₜ = ω + α · ε²ₜ₋₁ + β · σ²ₜ₋₁')
add_paragraph(
    'donde μ es la media condicional de los retornos, σₜ es la desviación típica condicional en el '
    'día t, zₜ sigue una distribución t-Student estandarizada con ν grados de libertad, ω es una '
    'constante positiva, α mide la reacción a los shocks (efecto ARCH) y β mide la persistencia '
    'de la volatilidad (efecto GARCH).'
)
add_paragraph(
    'La elección de innovaciones t-Student (en lugar de gaussianas) se justifica por los resultados '
    'del test de Jarque-Bera del EDA, que documentaron una curtosis de 16 en los log-retornos del IBEX 35. '
    'La distribución t-Student permite modelar estas colas pesadas a través del parámetro ν.'
)
add_paragraph('Parámetros estimados:', bold=True)
add_table_img(
    ['Parámetro', 'Valor', 'Interpretación'],
    [
        ['ω (constante)', '1,12e-06', 'Nivel base de varianza'],
        ['α (efecto ARCH)', '0,1001', 'Reacción moderada a shocks'],
        ['β (efecto GARCH)', '0,8722', 'Alta persistencia'],
        ['α + β', '0,9722', 'Estacionario (< 1)'],
        ['ν (grados de libertad)', '6,35', 'Colas pesadas confirmadas'],
    ],
    'tab_2_6_garch_params.png',
    caption='Tabla 7. Parámetros estimados del modelo GARCH(1,1) con t-Student.',
    col_widths=[0.25, 0.2, 0.45]
)
add_paragraph(
    'El valor α + β = 0,9722 confirma la estacionariedad del proceso (< 1) pero indica una persistencia '
    'muy alta: la vida media de un shock de volatilidad es −1/ln(0,9722) ≈ 35 días de negociación. '
    'Los grados de libertad ν = 6,35 corresponden a colas significativamente más pesadas que la normal '
    '(ν → ∞ equivale a normalidad).'
)
add_paragraph('Resultados en el conjunto de test:', bold=True)
add_table_img(
    ['Métrica', 'Valor'],
    [
        ['RMSE', '0,1029'],
        ['MAE', '0,0986'],
        ['R²', '−3,8375'],
        ['QLIKE', '—'],
    ],
    'tab_2_6_resultados_garch.png',
    caption='Tabla 8. Resultados del Modelo B (GARCH) en test.',
    col_widths=[0.4, 0.4]
)
add_paragraph(
    'El R² negativo del GARCH merece una explicación detallada, ya que no indica necesariamente que el '
    'modelo sea inútil. El GARCH modela la varianza condicional diaria (σ²ₜ) del proceso de retornos, '
    'que es conceptualmente distinta de la volatilidad realizada a 21 días (vol_hist_21d) que sirve como '
    'variable objetivo. La σₜ del GARCH mide la volatilidad instantánea un paso adelante, mientras que '
    'vol_hist_21d es una media móvil de 21 días de la volatilidad realizada. Esta discrepancia de escala '
    'temporal genera un sesgo sistemático que infla el RMSE y produce un R² negativo.',
    italic=True
)
add_paragraph(
    'A pesar de estas métricas aparentemente desfavorables, el GARCH cumple una función esencial en el '
    'análisis: captura la dinámica autorregresiva de la varianza condicional y proporciona estimaciones '
    'de los parámetros α y β que caracterizan la persistencia de la volatilidad del IBEX 35.'
)
add_figure('17_garch_resultados.png',
           'Figura 20. Modelo B (GARCH): volatilidad condicional estimada vs. volatilidad realizada en test.')
add_figure('17b_garch_diagnostico.png',
           'Figura 21. Modelo B (GARCH): diagnóstico de residuos estandarizados.')

# 2.7
doc.add_heading('2.7. Modelo C: XGBoost', level=2)
add_paragraph(
    'El modelo XGBoost (eXtreme Gradient Boosting) se entrena con las mismas 15 features que el OLS '
    '(3 HAR + 12 macro), pero puede capturar relaciones no lineales e interacciones entre variables. '
    'La configuración del modelo incluye:'
)
add_paragraph('Hiperparámetros:', bold=True)
add_table_img(
    ['Hiperparámetro', 'Valor', 'Justificación'],
    [
        ['n_estimators', '100', 'Seleccionado por CV temporal'],
        ['max_depth', '6', 'Control de complejidad'],
        ['learning_rate', '0,1', 'Tasa estándar con regularización'],
        ['subsample', '0,8', 'Reducción de varianza'],
        ['colsample_bytree', '0,8', 'Decorrelación de árboles'],
        ['reg_alpha (L1)', '0,1', 'Regularización Lasso'],
        ['reg_lambda (L2)', '1,0', 'Regularización Ridge'],
    ],
    'tab_2_7_xgb_params.png',
    caption='Tabla 9. Hiperparámetros del modelo XGBoost.',
    col_widths=[0.3, 0.15, 0.45]
)
add_paragraph(
    'La selección del número óptimo de árboles (n_estimators) se realiza mediante validación cruzada '
    'temporal con TimeSeriesSplit de scikit-learn (5 folds), que preserva el orden cronológico de las '
    'observaciones y evita data leakage. El número óptimo resultante es 100 árboles.'
)
add_paragraph('Resultados en el conjunto de test:', bold=True)
add_table_img(
    ['Métrica', 'Valor'],
    [
        ['RMSE', '0,0082'],
        ['MAE', '0,0052'],
        ['R²', '0,9691'],
        ['QLIKE', '−1,8419'],
    ],
    'tab_2_7_resultados_xgb.png',
    caption='Tabla 10. Resultados del Modelo C (XGBoost) en test.',
    col_widths=[0.4, 0.4]
)
add_paragraph(
    'XGBoost obtiene el mejor RMSE (0,0082) y el mejor R² (0,9691) de los cuatro modelos. La mejora '
    'respecto a la regresión simple es modesta en términos absolutos (0,0001 en RMSE), pero el hecho '
    'de que lo consiga con 15 features (frente a 1) sugiere que captura información no lineal de las '
    'variables macroeconómicas que el OLS no puede explotar.'
)

add_paragraph('Importancia de features por SHAP:', bold=True)
add_paragraph(
    'Los valores SHAP (Lundberg y Lee, 2017) revelan la importancia relativa de cada feature en las '
    'predicciones del XGBoost:'
)
add_table_img(
    ['Feature', 'Mean |SHAP|', '% Total'],
    [
        ['vol_lag1', '0,0547', '82,4%'],
        ['vol_lag5', '0,0080', '12,1%'],
        ['vix', '0,0026', '3,9%'],
        ['vol_lag21', '0,0003', '0,5%'],
        ['bono_es_10y', '0,0002', '0,3%'],
        ['tasa_paro', '0,0002', '0,3%'],
        ['Resto (9 vars.)', '0,0003', '0,5%'],
    ],
    'tab_2_7_shap.png',
    caption='Tabla 11. Importancia de features por SHAP en XGBoost (top 6 + resto).',
    col_widths=[0.3, 0.25, 0.2]
)
add_paragraph(
    'La dominancia de vol_lag1 (82,4% de la importancia total) confirma que la persistencia de la '
    'volatilidad es el mecanismo predictivo fundamental. Sin embargo, el VIX emerge como la variable '
    'macroeconómica más informativa (3,9%), seguida del bono español a 10 años y la tasa de paro. '
    'Esto sugiere que el canal de contagio global (VIX) y el riesgo soberano (bono) proporcionan '
    'información complementaria que XGBoost explota de forma no lineal.'
)
add_figure('18_xgboost_resultados.png',
           'Figura 22. Modelo C (XGBoost): predicción vs. valores reales en el conjunto de test.')
add_figure('18b_xgboost_diagnostico.png',
           'Figura 23. Modelo C (XGBoost): diagnóstico de residuos y valores SHAP.')

# 2.8
doc.add_heading('2.8. Comparación de modelos', level=2)
add_paragraph(
    'La siguiente tabla resume las métricas de los cuatro modelos en el conjunto de test, con '
    'sombreado para identificar el mejor resultado en cada métrica:'
)
add_table_img(
    ['Modelo', 'RMSE', 'MAE', 'R²', 'QLIKE'],
    [
        ['A0: Simple', '0,0083', '0,0048 *', '0,9686', '−1,8419'],
        ['A: OLS HAR+Macro', '0,0084', '0,0061', '0,9677', '−1,8416'],
        ['B: GARCH(1,1)', '0,1029', '0,0986', '−3,8375', '—'],
        ['C: XGBoost', '0,0082 *', '0,0052', '0,9691 *', '−1,8419 *'],
    ],
    'tab_2_8_comparacion.png',
    caption='Tabla 12. Comparación de los cuatro modelos (* = mejor resultado en la métrica).',
    title='Comparación de modelos predictivos',
    col_widths=[0.28, 0.15, 0.15, 0.15, 0.15]
)
add_paragraph(
    'XGBoost obtiene los mejores resultados en RMSE, R² y QLIKE, mientras que la Regresión Simple '
    'obtiene el mejor MAE. El OLS con 15 features es inferior a la regresión simple en todas las métricas, '
    'confirmando que en el marco lineal las variables macro no aportan valor predictivo. El GARCH no es '
    'directamente comparable por la discrepancia de escala temporal ya explicada.'
)

add_paragraph('Test de Diebold-Mariano (DM):', bold=True)
add_paragraph(
    'Para determinar si las diferencias en capacidad predictiva son estadísticamente significativas, '
    'se aplica el test de Diebold y Mariano (1995) a cada par de modelos (excluyendo GARCH). El test '
    'evalúa la hipótesis nula de que ambos modelos tienen la misma precisión predictiva:'
)
add_table_img(
    ['Par de modelos', 'DM stat', 'p-valor', 'Significativo'],
    [
        ['Simple vs. OLS', '−0,693', '0,489', 'NO'],
        ['Simple vs. XGBoost', '0,315', '0,753', 'NO'],
        ['OLS vs. XGBoost', '0,698', '0,485', 'NO'],
        ['Simple vs. GARCH', '−27,21', '≈ 0,000', 'SÍ'],
        ['OLS vs. GARCH', '−27,18', '≈ 0,000', 'SÍ'],
        ['XGBoost vs. GARCH', '−27,23', '≈ 0,000', 'SÍ'],
    ],
    'tab_2_8_dm.png',
    caption='Tabla 13. Resultados del test de Diebold-Mariano entre pares de modelos.',
    highlight_col=3,
    highlight_vals={'SÍ': '#D4EDDA', 'NO': '#FFF3CD'},
    col_widths=[0.3, 0.15, 0.15, 0.18]
)
add_paragraph(
    'No existe diferencia significativa entre Simple, OLS y XGBoost (p > 0,49 en los tres pares). '
    'Esto implica que, a pesar de las diferencias numéricas en RMSE, ningún modelo es estadísticamente '
    'superior a los otros para predecir la volatilidad del IBEX 35. El GARCH es significativamente peor '
    '(p ≈ 0), pero esto se debe a la discrepancia de escala temporal y no a una inferioridad intrínseca '
    'del modelo.',
    bold=True
)
add_figure('19_comparacion_modelos.png',
           'Figura 24. Comparación visual de los cuatro modelos: predicción vs. real, scatter y métricas.')

# 2.9
doc.add_heading('2.9. Valor añadido de las variables macroeconómicas', level=2)
add_paragraph(
    'Una pregunta central de este trabajo es si las variables macroeconómicas aportan información '
    'predictiva más allá de la propia historia de la volatilidad. Para responderla, se compara '
    'cada modelo con y sin variables macro:'
)
add_table_img(
    ['Modelo', 'Solo HAR (RMSE)', 'HAR + Macro (RMSE)', 'Δ RMSE', 'Efecto'],
    [
        ['OLS', '0,0078', '0,0084', '+7,2%', 'Empeora'],
        ['XGBoost', '0,0095', '0,0082', '−14,1%', 'Mejora'],
    ],
    'tab_2_9_macro.png',
    caption='Tabla 14. Impacto de las variables macroeconómicas en la capacidad predictiva.',
    highlight_col=4,
    highlight_vals={'Empeora': '#FDDCDC', 'Mejora': '#D4EDDA'},
    col_widths=[0.15, 0.2, 0.22, 0.12, 0.12]
)
add_paragraph(
    'Este resultado es uno de los hallazgos más reveladores del trabajo. El OLS empeora un 7,2% al añadir '
    'variables macro, lo que indica sobreajuste (overfitting): con 15 features y una relación lineal, '
    'el modelo captura ruido en lugar de señal. En cambio, XGBoost mejora un 14,1%, demostrando que '
    'la información macroeconómica sí contiene señal predictiva, pero su relación con la volatilidad '
    'es no lineal y requiere un modelo flexible para ser explotada.'
)
add_paragraph(
    'Las tres variables macro más importantes según SHAP en el modelo XGBoost son: VIX (canal de '
    'contagio global), bono español a 10 años (riesgo soberano) y tasa de paro (ciclo económico real). '
    'Estas tres variables representan los tres canales de transmisión macroeconómica hacia la '
    'volatilidad bursátil identificados en el marco teórico.'
)
add_figure('20_valor_anadido_macro.png',
           'Figura 25. Valor añadido de las variables macroeconómicas: comparación HAR solo vs. HAR + Macro para OLS y XGBoost.')

# 2.10
doc.add_heading('2.10. Diagnóstico de residuos', level=2)
add_paragraph(
    'El diagnóstico de residuos es un paso fundamental para evaluar la adecuación de los modelos y '
    'detectar posibles violaciones de los supuestos estadísticos. Para cada modelo se examinan cuatro '
    'dimensiones:'
)
add_paragraph('Normalidad de residuos:', bold=True)
add_paragraph(
    'Ningún modelo produce residuos normales (test Jarque-Bera significativo en todos los casos). '
    'Esto es esperable en datos financieros y no invalida las estimaciones puntuales, pero aconseja '
    'cautela al interpretar intervalos de confianza basados en normalidad.'
)
add_paragraph('Autocorrelación de residuos:', bold=True)
add_paragraph(
    'Los modelos de regresión (Simple, OLS, XGBoost) presentan autocorrelación residual significativa '
    'en los primeros lags (test Ljung-Box significativo), lo que indica que queda estructura temporal '
    'por capturar. Esto es una consecuencia de la memoria larga de la volatilidad: incluso con el '
    'esquema HAR de tres escalas, persiste autocorrelación en lags intermedios.'
)
add_paragraph('Homocedasticidad:', bold=True)
add_paragraph(
    'Los residuos de todos los modelos exhiben heterocedasticidad condicional: los errores son mayores '
    'en períodos de alta volatilidad y menores en períodos tranquilos. Este fenómeno de "volatilidad '
    'de la volatilidad" es intrínseco a las series financieras.'
)
add_paragraph('Patrón temporal de errores:', bold=True)
add_paragraph(
    'Los errores más grandes se concentran en transiciones de régimen (entrada y salida de crisis), '
    'donde la volatilidad cambia bruscamente. Todos los modelos subestiman estos picos transitorios.'
)

# ════════════════════════════════════════════════════════════════════
# 3. CAPÍTULO 2: EVENT STUDY
# ════════════════════════════════════════════════════════════════════
doc.add_heading('3. Capítulo 2: Event Study — Impacto de Publicaciones Macroeconómicas', level=1)

# 3.1
doc.add_heading('3.1. Motivación y metodología', level=2)
add_paragraph(
    'Los resultados del Capítulo 1 sugieren que los niveles de las variables macroeconómicas tienen '
    'un poder predictivo limitado sobre la volatilidad futura del IBEX 35 (la persistencia de la '
    'propia volatilidad domina). Sin embargo, esto no descarta que las publicaciones macroeconómicas '
    '—es decir, los momentos en los que nueva información llega al mercado— generen reacciones de '
    'corto plazo en la volatilidad.'
)
add_paragraph(
    'Esta hipótesis se fundamenta en la Hipótesis de los Mercados Eficientes (Fama, 1970): si los '
    'mercados son eficientes, los precios deben ajustarse rápidamente ante nueva información, y este '
    'ajuste se manifestaría como un aumento transitorio de la volatilidad en las horas posteriores a '
    'la publicación.'
)
add_paragraph(
    'Para contrastar esta hipótesis se diseña un estudio de eventos (event study) con la siguiente '
    'metodología:'
)

add_paragraph('Definición de evento:', bold=True)
add_paragraph(
    'Un evento se define como un día de negociación en el que la variable macroeconómica forward-filled '
    'cambia de valor (Δmacro ≠ 0). Este criterio captura las fechas de publicación oficial de cada '
    'indicador sin depender de un calendario externo.'
)

add_paragraph('Benchmark de referencia:', bold=True)
add_paragraph(
    'Para cada evento, se calcula la volatilidad intradía de referencia como la mediana de |log_ret| '
    'en la ventana de pre-evento [-10, -1] (los 10 días de negociación anteriores). Este benchmark '
    'refleja el nivel "normal" de volatilidad antes del evento.'
)

add_paragraph('Volatilidad anormal:', bold=True)
add_paragraph('    AR_t = |log_ret|_evento − mediana(|log_ret|_{[-10,-1]})')
add_paragraph(
    'Un valor positivo de AR_t indica que la volatilidad en el día del evento fue superior a lo normal, '
    'sugiriendo una reacción del mercado a la publicación.'
)

# 3.2
doc.add_heading('3.2. Identificación de eventos', level=2)
add_paragraph(
    'Se identifican un total de 1.058 eventos de publicación macroeconómica distribuidos en 7 '
    'indicadores. La tabla siguiente detalla el número de eventos por variable:'
)
add_table_img(
    ['Indicador', 'Eventos', 'Frecuencia aprox.', 'Período'],
    [
        ['ipc_yoy', '194', 'Mensual', '2008 – 2025'],
        ['euribor_3m', '185', 'Mensual', '2005 – 2025'],
        ['tipo_dfr (BCE)', '179', 'Irregular', '2009 – 2025'],
        ['bono_es_10y', '167', '~ Mensual', '2005 – 2025'],
        ['tasa_paro', '126', 'Trimestral', '2005 – 2025'],
        ['pib_yoy', '115', 'Trimestral', '2005 – 2025'],
        ['ipi_yoy', '92', 'Mensual', '2005 – 2025'],
        ['TOTAL', '1.058', '—', '—'],
    ],
    'tab_3_2_eventos.png',
    caption='Tabla 15. Número de eventos de publicación macroeconómica identificados por indicador.',
    col_widths=[0.2, 0.15, 0.25, 0.25]
)

# 3.3
doc.add_heading('3.3. Resultados agregados', level=2)
add_paragraph(
    'El primer análisis compara la distribución de volatilidad intradía (|log_ret|) entre los días '
    'con evento macroeconómico y los días sin evento, utilizando el test no paramétrico de Mann-Whitney U:'
)
add_table_img(
    ['Grupo', 'Mediana |log_ret|', 'N'],
    [
        ['Días con evento', '0,01329', '1.058'],
        ['Días sin evento', '0,01191', '2.397'],
        ['Diferencia', '+0,00138 (+11,6%)', '—'],
    ],
    'tab_3_3_agregados.png',
    caption='Tabla 16. Comparación de volatilidad intradía entre días con y sin evento macroeconómico.',
    col_widths=[0.35, 0.3, 0.15]
)
add_paragraph(
    'Mann-Whitney U: p = 0,0000 (significativo al 1%).',
    bold=True
)
add_paragraph(
    'La diferencia es altamente significativa: los días en los que se publica un indicador macroeconómico '
    'presentan una volatilidad intradía un 11,6% superior a los días sin publicación. Este resultado '
    'confirma la hipótesis de que las publicaciones macroeconómicas generan shocks de volatilidad en '
    'el IBEX 35.'
)
add_paragraph(
    'Además, el 61,9% de los eventos muestran volatilidad anormal positiva (AR_t > 0), lo que indica '
    'que la reacción del mercado es predominantemente en dirección de mayor volatilidad.'
)
add_figure('20_event_study_panorama.png',
           'Figura 26. Event Study — Panorama general: distribución de volatilidad en días con y sin evento, '
           'y porcentaje de eventos con volatilidad anormal positiva.')

# 3.4
doc.add_heading('3.4. Resultados por indicador', level=2)
add_paragraph(
    'Para cada indicador macroeconómico se aplica el test de Wilcoxon de rangos signados, que evalúa '
    'si la mediana de la volatilidad anormal es significativamente distinta de cero:'
)
add_table_img(
    ['Indicador', 'Mediana AR', 'Wilcoxon p', 'Significativo', '% AR > 0'],
    [
        ['pib_yoy', '+0,0031', '0,0002', 'SÍ ***', '62,6%'],
        ['tasa_paro', '+0,0030', '0,0019', 'SÍ **', '60,3%'],
        ['ipc_yoy', '+0,0017', '0,0011', 'SÍ **', '59,3%'],
        ['ipi_yoy', '+0,0024', '0,0038', 'SÍ **', '58,7%'],
        ['euribor_3m', '+0,0012', '0,0107', 'SÍ *', '57,8%'],
        ['bono_es_10y', '+0,0011', '0,0213', 'SÍ *', '56,3%'],
        ['tipo_dfr (BCE)', '+0,0009', '0,0584', 'Límite', '54,2%'],
    ],
    'tab_3_4_wilcoxon.png',
    caption='Tabla 17. Test de Wilcoxon por indicador (*** p<0.001, ** p<0.01, * p<0.05).',
    highlight_col=3,
    highlight_vals={
        'SÍ ***': '#D4EDDA', 'SÍ **': '#D4EDDA', 'SÍ *': '#D4EDDA',
        'Límite': '#FFF3CD',
    },
    col_widths=[0.18, 0.15, 0.15, 0.18, 0.12]
)
add_paragraph(
    'Seis de los siete indicadores generan un impacto significativo en la volatilidad del IBEX 35. '
    'El PIB interanual tiene el mayor impacto (+0,0031 de volatilidad anormal mediana), seguido de '
    'la tasa de paro (+0,0030). Solo el tipo de referencia del BCE se queda en el límite de significación '
    '(p = 0,0584), probablemente porque las decisiones del BCE suelen estar ampliamente anticipadas por '
    'el mercado (forward guidance).'
)

add_paragraph('Test de Kruskal-Wallis entre indicadores:', bold=True)
add_paragraph(
    'Para evaluar si el tipo de publicación importa (no todos los indicadores tienen el mismo impacto), '
    'se aplica el test de Kruskal-Wallis a la volatilidad anormal entre los 7 indicadores:'
)
add_paragraph('    H = 10,71    p = 0,0047', bold=True)
add_paragraph(
    'El test rechaza la igualdad de distribuciones (p < 0,01), confirmando que no todas las publicaciones '
    'macroeconómicas generan el mismo impacto. Las publicaciones de actividad económica real (PIB, '
    'desempleo) generan shocks significativamente mayores que las de política monetaria (BCE).'
)
add_figure('21_event_study_detalle.png',
           'Figura 27. Event Study — Detalle por indicador: boxplots de volatilidad anormal, tests de Wilcoxon '
           'y barras de porcentaje de eventos positivos.')

# 3.5
doc.add_heading('3.5. Efecto asimétrico del signo', level=2)
add_paragraph(
    'Un resultado particularmente interesante es el efecto asimétrico del signo de la publicación. '
    'No todas las "sorpresas" generan la misma reacción: el mercado puede reaccionar de forma distinta '
    'ante buenas y malas noticias macroeconómicas.'
)
add_paragraph('Hallazgos principales:', bold=True)

p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Tipo BCE (tipo_dfr): ')
run_b.bold = True
p.add_run(
    'las bajadas de tipos generan más volatilidad que las subidas. Esto puede parecer contraintuitivo '
    '(las bajadas suelen ser "buenas noticias"), pero refleja que las bajadas a menudo se producen en '
    'contextos de crisis que amplifican la incertidumbre.'
)

p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Tasa de paro: ')
run_b.bold = True
p.add_run(
    'los aumentos del desempleo generan más volatilidad que las disminuciones. La asimetría es '
    'consistente con la evidencia de que los mercados reaccionan más a las malas noticias que a las '
    'buenas (efecto leverage en finanzas).'
)

p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('IPC interanual: ')
run_b.bold = True
p.add_run(
    'las caídas de la inflación generan más volatilidad que los aumentos. Esto puede reflejar que '
    'las caídas de inflación se asocian con riesgo de deflación, un escenario que los mercados '
    'perciben como más peligroso.'
)

add_figure('22_event_study_bce_ipc.png',
           'Figura 28. Efecto asimétrico del signo: impacto en volatilidad de subidas vs. bajadas para BCE e IPC.')

# 3.6
doc.add_heading('3.6. Persistencia y magnitud', level=2)
add_paragraph(
    'El último análisis evalúa dos dimensiones adicionales del impacto de las publicaciones '
    'macroeconómicas: la persistencia temporal del shock y la relación entre la magnitud del cambio '
    'y el impacto en volatilidad.'
)

add_paragraph('Persistencia temporal:', bold=True)
add_paragraph(
    'La volatilidad anormal no se disipa completamente el día del evento. Se observa que la volatilidad '
    'permanece por encima del benchmark de pre-evento tanto en t+1 (primer día post-publicación) como '
    'en t+2 (segundo día), aunque la magnitud del exceso disminuye progresivamente. Esto sugiere que '
    'el mercado necesita varios días para absorber completamente la nueva información macroeconómica.'
)

add_paragraph('Relación magnitud-impacto:', bold=True)
add_paragraph(
    'Los scatter plots de |Δmacro| (magnitud absoluta del cambio en la variable macroeconómica) '
    'contra la volatilidad anormal revelan una relación positiva pero débil y ruidosa. Cambios '
    'macroeconómicos grandes tienden a generar mayor volatilidad anormal, pero la relación no es '
    'determinista: factores como el contexto de mercado, la anticipación del dato y la sorpresa '
    'respecto al consenso modulan la reacción.'
)
add_figure('23_event_study_magnitud.png',
           'Figura 29. Persistencia del shock post-evento y relación magnitud del cambio macroeconómico '
           'vs. volatilidad anormal.')

# ════════════════════════════════════════════════════════════════════
# 4. CONCLUSIONES
# ════════════════════════════════════════════════════════════════════
doc.add_heading('4. Conclusiones del Análisis del Dato', level=1)

add_paragraph(
    'El análisis del dato de este trabajo se ha articulado en dos capítulos complementarios —modelos '
    'predictivos y event study— que, tomados conjuntamente, proporcionan una visión integral del papel '
    'de las variables macroeconómicas en la volatilidad del IBEX 35. Las conclusiones principales son:'
)

add_paragraph('1. La persistencia de la volatilidad domina la predicción.', bold=True)
add_paragraph(
    'El resultado más robusto del Capítulo 1 es que un solo lag de volatilidad (vol_lag1) explica '
    'el 96,9% de la varianza de la volatilidad futura (R² = 0,969). Los tres modelos alternativos '
    '(OLS completo, GARCH y XGBoost) no mejoran significativamente esta cifra según el test de '
    'Diebold-Mariano (p > 0,49). Este hallazgo tiene una implicación práctica directa: para la '
    'predicción a un día, la volatilidad reciente es el input más valioso.'
)

add_paragraph('2. XGBoost es el mejor modelo pero sin significación estadística.', bold=True)
add_paragraph(
    'XGBoost obtiene el mejor RMSE (0,0082), R² (0,969) y QLIKE (−1,842), pero el test de '
    'Diebold-Mariano no encuentra diferencia significativa con la regresión simple (p = 0,753). '
    'Esto refleja la dificultad inherente de mejorar sobre un proceso tan fuertemente autorregresivo '
    'como la volatilidad.'
)

add_paragraph('3. Las variables macro aportan valor de forma no lineal.', bold=True)
add_paragraph(
    'Añadir variables macroeconómicas al OLS empeora el RMSE en un 7,2% (sobreajuste lineal), '
    'pero mejora el RMSE del XGBoost en un 14,1%. Esto demuestra que la información macroeconómica '
    'contiene señal predictiva, pero su relación con la volatilidad es no lineal y requiere un '
    'modelo flexible (como XGBoost) para ser explotada. Las variables más informativas son el VIX '
    '(contagio global), el bono español a 10 años (riesgo soberano) y la tasa de paro (ciclo económico).'
)

add_paragraph('4. Las publicaciones macroeconómicas generan shocks significativos.', bold=True)
add_paragraph(
    'El event study del Capítulo 2 demuestra que los días de publicación macroeconómica presentan una '
    'volatilidad intradía un 11,6% superior a los días normales (Mann-Whitney U, p ≈ 0). Este resultado '
    'reconcilia la aparente contradicción entre el bajo poder predictivo de los niveles macro (Capítulo 1) '
    'y su relevancia económica: los niveles no predicen, pero las publicaciones (la llegada de nueva '
    'información) sí generan reacciones inmediatas.'
)

add_paragraph('5. El PIB y la tasa de paro tienen el mayor impacto.', bold=True)
add_paragraph(
    'Entre los 7 indicadores analizados, las publicaciones de PIB interanual (+0,0031 de volatilidad '
    'anormal mediana) y tasa de paro (+0,0030) generan los shocks más fuertes. Esto es consistente con '
    'la relevancia de la actividad económica real para un índice tan bancario como el IBEX 35. Solo el '
    'tipo del BCE no alcanza significación estadística (p = 0,058), reflejando la eficacia del forward '
    'guidance del BCE para anticipar sus decisiones.'
)

add_paragraph('6. El signo de la publicación importa de forma asimétrica.', bold=True)
add_paragraph(
    'Las bajadas de tipos BCE generan más volatilidad que las subidas, los aumentos de paro más que '
    'las bajadas, y las caídas de inflación más que las subidas. Esta asimetría es consistente con la '
    'evidencia de que los mercados financieros reaccionan más intensamente ante señales de deterioro '
    'económico que ante señales de mejora.'
)

add_paragraph('7. Limitaciones y trabajo futuro.', bold=True)
add_paragraph(
    'Las principales limitaciones de este trabajo incluyen: (a) el uso de volatilidad realizada a 21 días '
    'como proxy, que introduce autocorrelación mecánica; (b) la ausencia de datos de consenso para '
    'medir la sorpresa macroeconómica; (c) la limitación a predicción a un día (horizontes más largos '
    'podrían beneficiarse más de las variables macro); y (d) la no consideración de efectos asimétricos '
    'de los retornos en el GARCH (modelos E-GARCH o GJR-GARCH). Líneas futuras de investigación incluyen '
    'la incorporación de datos de consenso de analistas, la extensión a horizontes de predicción más '
    'largos, el uso de modelos de deep learning (LSTM, Transformers) y la inclusión de datos de texto '
    '(sentiment analysis de noticias financieras).'
)

# ════════════════════════════════════════════════════════════════════
# 5. RESUMEN DE FIGURAS
# ════════════════════════════════════════════════════════════════════
doc.add_heading('5. Resumen de figuras', level=1)
add_paragraph(
    'La siguiente tabla recoge las 15 figuras incluidas en este documento, con su numeración original '
    'del pipeline de análisis, título descriptivo y sección de referencia.'
)
add_table_img(
    ['Fig.', 'Título', 'Sección'],
    [
        ['15', 'Split temporal train/test', '2.1'],
        ['16', 'Modelo A0 (Simple): predicción vs. real', '2.4'],
        ['17', 'Modelo A0 (Simple): diagnóstico de residuos', '2.4'],
        ['18', 'Modelo A (OLS): predicción vs. real', '2.5'],
        ['19', 'Modelo A (OLS): diagnóstico de residuos', '2.5'],
        ['20', 'Modelo B (GARCH): volatilidad condicional vs. realizada', '2.6'],
        ['21', 'Modelo B (GARCH): diagnóstico de residuos', '2.6'],
        ['22', 'Modelo C (XGBoost): predicción vs. real', '2.7'],
        ['23', 'Modelo C (XGBoost): diagnóstico y SHAP', '2.7'],
        ['24', 'Comparación de los cuatro modelos', '2.8'],
        ['25', 'Valor añadido de las variables macro', '2.9'],
        ['26', 'Event Study: panorama general', '3.3'],
        ['27', 'Event Study: detalle por indicador', '3.4'],
        ['28', 'Efecto asimétrico del signo (BCE e IPC)', '3.5'],
        ['29', 'Persistencia y magnitud del shock', '3.6'],
    ],
    'tab_5_resumen_figuras.png',
    caption='Tabla 18. Resumen de las 15 figuras del documento de Análisis del Dato.',
    col_widths=[0.06, 0.68, 0.1]
)

# ── GUARDAR ────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Documento guardado en:\n  {OUT}")
print(f"Tablas renderizadas en:\n  {TABS}")
print(f"Total tablas imagen: {len(os.listdir(TABS))}")
