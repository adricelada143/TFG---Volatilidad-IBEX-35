"""
Genera el archivo Word de la memoria de Ingeniería del Dato
con todas las figuras y tablas estilizadas como imágenes.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import tempfile
import shutil

# ── Rutas ──────────────────────────────────────────────────────────
BASE = "/Users/adriancelada/Library/Mobile Documents/com~apple~CloudDocs/UFV/UNIVERSIDAD FRANCISCO DE VITORIA/4º/TFG"
FIGS = os.path.join(BASE, "proyecto/outputs/figuras")
TABS = os.path.join(BASE, "proyecto/outputs/tablas_img")
OUT  = os.path.join(BASE, "Ingenieria del Dato/Memoria_Ingenieria_del_Dato.docx")

os.makedirs(TABS, exist_ok=True)

doc = Document()

# ── Estilos ────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0, 51, 102)
    if level == 1:
        hs.font.size = Pt(16)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(13)
        hs.font.bold = True
    else:
        hs.font.size = Pt(11)
        hs.font.bold = True
        hs.font.italic = True


# ── Funciones auxiliares ───────────────────────────────────────────
def add_paragraph(text, bold=False, italic=False, style_name='Normal'):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_figure(filename, caption, width=Inches(5.8)):
    path = os.path.join(FIGS, filename)
    if not os.path.exists(path):
        add_paragraph(f"[Figura no encontrada: {filename}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(80, 80, 80)


def render_table_image(headers, rows, filename, title=None, col_widths=None,
                       highlight_col=None, highlight_vals=None):
    """
    Renderiza una tabla como imagen PNG estilizada.
    - headers: lista de strings
    - rows: lista de listas de strings
    - highlight_col: índice de columna para colorear celdas condicionalmente
    - highlight_vals: dict {valor: color_hex} para resaltar celdas
    """
    n_rows = len(rows)
    n_cols = len(headers)

    # Calcular anchos automáticos basados en contenido
    if col_widths is None:
        max_lens = []
        for c in range(n_cols):
            max_len = len(str(headers[c]))
            for r in rows:
                max_len = max(max_len, len(str(r[c])))
            max_lens.append(max_len)
        total = sum(max_lens)
        col_widths = [m / total for m in max_lens]

    fig_width = max(6, min(12, n_cols * 1.8))
    row_height = 0.38
    fig_height = (n_rows + 1) * row_height + (0.5 if title else 0.15) + 0.15

    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, n_rows + 1)
    ax.axis('off')

    # Colores
    HEADER_BG = '#003366'
    HEADER_FG = '#FFFFFF'
    ROW_EVEN  = '#F7F9FC'
    ROW_ODD   = '#FFFFFF'
    BORDER    = '#D0D5DD'
    TEXT_COLOR = '#1A1A2E'

    # Título
    if title:
        fig.suptitle(title, fontsize=10, fontweight='bold', color='#003366',
                     y=1 - 0.08 / fig_height, ha='center')

    # Dibujar celdas
    for row_idx in range(n_rows + 1):  # +1 para header
        y = n_rows - row_idx  # invertir para que header arriba

        for col_idx in range(n_cols):
            x_start = sum(col_widths[:col_idx])
            w = col_widths[col_idx]

            # Background
            if row_idx == 0:
                bg = HEADER_BG
                fg = HEADER_FG
                fw = 'bold'
                fs = 8.5
            else:
                bg = ROW_EVEN if row_idx % 2 == 0 else ROW_ODD
                fg = TEXT_COLOR
                fw = 'normal'
                fs = 8

                # Highlight condicional
                if highlight_col is not None and highlight_vals and col_idx == highlight_col:
                    val = str(rows[row_idx - 1][col_idx]).strip()
                    if val in highlight_vals:
                        bg = highlight_vals[val]

            rect = plt.Rectangle((x_start, y), w, 1,
                                 facecolor=bg, edgecolor=BORDER, linewidth=0.5)
            ax.add_patch(rect)

            # Texto
            text = headers[col_idx] if row_idx == 0 else str(rows[row_idx - 1][col_idx])
            ax.text(x_start + w / 2, y + 0.5, text,
                    ha='center', va='center', fontsize=fs,
                    color=fg, fontweight=fw, fontfamily='sans-serif')

    plt.tight_layout(pad=0.2)
    path = os.path.join(TABS, filename)
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(fig)
    return path


def add_table_img(headers, rows, filename, caption=None, title=None,
                  width=Inches(5.5), highlight_col=None, highlight_vals=None,
                  col_widths=None):
    """Renderiza tabla como imagen y la inserta en el doc."""
    path = render_table_image(headers, rows, filename, title=title,
                              highlight_col=highlight_col,
                              highlight_vals=highlight_vals,
                              col_widths=col_widths)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(80, 80, 80)


# ════════════════════════════════════════════════════════════════════
#  CONTENIDO
# ════════════════════════════════════════════════════════════════════

# ── TÍTULO ─────────────────────────────────────────────────────────
title = doc.add_heading('Ingeniería del Dato', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)
doc.add_paragraph()

# ── 1. INTRODUCCIÓN ────────────────────────────────────────────────
doc.add_heading('1. Introducción', level=1)
add_paragraph(
    'La fase de ingeniería del dato constituye el pilar fundamental sobre el que se construye '
    'todo el análisis posterior de este trabajo. Su objetivo es construir un dataset maestro que '
    'integre, en una única tabla de frecuencia diaria, los precios de las 35 empresas del IBEX 35 '
    'junto con 20 variables macroeconómicas de distinta naturaleza y frecuencia temporal. Este proceso '
    'abarca la extracción de datos desde múltiples fuentes heterogéneas, su transformación y limpieza, '
    'la alineación temporal al calendario bursátil español y un análisis exploratorio exhaustivo que '
    'justifica las decisiones de diseño adoptadas para la fase de modelización.'
)
add_paragraph(
    'El pipeline completo se implementa en cuatro notebooks secuenciales de Python y culmina en una '
    'base de datos SQLite (tfg.db) con 8 tablas y más de 325.000 registros.'
)

# ── 2. HERRAMIENTAS ────────────────────────────────────────────────
doc.add_heading('2. Herramientas tecnológicas', level=1)
add_paragraph(
    'El entorno de desarrollo utilizado es Python 3 ejecutado sobre Jupyter Notebook, lo que permite '
    'combinar código, visualizaciones y documentación en un flujo de trabajo reproducible. Las principales '
    'librerías empleadas son:'
)
tools = [
    ('pandas y NumPy', 'manipulación de DataFrames, cálculos vectorizados, gestión de fechas y operaciones de ventana móvil.'),
    ('SQLAlchemy y sqlite3', 'conexión y escritura en la base de datos relacional SQLite, elegida por su ligereza y portabilidad (no requiere servidor).'),
    ('matplotlib y seaborn', 'visualización estadística con configuración personalizada (DPI 120, estilo whitegrid, paleta deep).'),
    ('scipy.stats', 'tests de normalidad (Jarque-Bera), detección de outliers (Z-score), tests no paramétricos (Mann-Whitney U, Kruskal-Wallis).'),
    ('statsmodels', 'test de estacionariedad Augmented Dickey-Fuller (ADF), funciones de autocorrelación (ACF/PACF) y cross-correlación (CCF).'),
]
for name, desc in tools:
    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run(f'{name}: ')
    run_bold.bold = True
    p.add_run(desc)
add_paragraph(
    'La base de datos SQLite fue elegida frente a alternativas como PostgreSQL o CSV planos porque '
    'ofrece un equilibrio óptimo entre rendimiento, portabilidad y facilidad de consulta SQL, sin '
    'requerir la instalación ni administración de un servidor externo.'
)

# ── 3. ORIGEN DE LOS DATOS ────────────────────────────────────────
doc.add_heading('3. Origen de los datos', level=1)
add_paragraph(
    'Los datos de este trabajo provienen de siete fuentes distintas, lo que refleja la naturaleza '
    'heterogénea de la información financiera y macroeconómica. Cada fuente presenta formatos, '
    'frecuencias y convenciones diferentes, lo que exige un proceso de extracción específico para cada una.'
)

# 3.1
doc.add_heading('3.1. Reuters Eikon — Precios de empresas', level=2)
add_paragraph(
    'La fuente principal de datos bursátiles es la plataforma Reuters Eikon (Refinitiv), terminal '
    'profesional de información financiera. Se extrajeron dos tipos de datos:'
)
p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Fichero de referencia ')
run_b.bold = True
p.add_run(
    '(GridExport_November_5_2025_21_27_3.xlsx): exportación de la herramienta Grid con el listado '
    'de las 35 empresas del IBEX 35, incluyendo su RIC (Reuters Instrument Code), nombre completo, '
    'ticker de mercado, sector GICS (Global Industry Classification Standard) e ISIN.'
)
p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('35 archivos de precios históricos ')
run_b.bold = True
p.add_run(
    '(uno por empresa, formato .xlsx): exportación de la sección Price History de Reuters Eikon, '
    'con datos diarios OHLCV (Open, High, Low, Close, Volume) desde enero de 2005 hasta octubre de 2025.'
)
add_paragraph(
    'Los archivos de Reuters Eikon presentan una particularidad técnica: incluyen metadatos y '
    'cabeceras institucionales en las primeras ~30 filas del archivo Excel, antes de que comiencen '
    'los datos reales. Esto requiere una detección automática de la fila de cabecera, que se resuelve '
    'buscando la cadena "Exchange Date" en cada fila del archivo.'
)

# 3.2
doc.add_heading('3.2. Reuters Eikon — Variables macroeconómicas', level=2)
add_paragraph(
    'De Reuters Eikon también se obtienen varias series macroeconómicas de frecuencia mensual y trimestral:'
)
add_table_img(
    ['Variable', 'Descripción', 'Frecuencia'],
    [
        ['pib_yoy', 'PIB de España, variación interanual (%)', 'Trimestral'],
        ['tasa_paro', 'Tasa de desempleo de España (%)', 'Trimestral'],
        ['pmi', "Purchasing Managers' Index de España", 'Mensual'],
        ['ipc_yoy', 'Índice de Precios al Consumo, var. interanual (%)', 'Mensual'],
        ['gas_ttf', 'Gas Natural TTF, precio ref. europeo (EUR/MWh)', 'Diario'],
        ['eur_usd', 'Tipo de cambio Euro/Dólar', 'Diario'],
    ],
    'tab_3_2_reuters_macro.png',
    caption='Tabla 1. Variables macroeconómicas obtenidas de Reuters Eikon.',
    col_widths=[0.15, 0.6, 0.15]
)
add_paragraph(
    'Estos archivos utilizan un formato de hoja de cálculo con los datos en la pestaña "First Release Data", '
    'donde los períodos se expresan como cadenas de texto ("Q1 2005" para trimestrales, "Jan 2005" para mensuales) '
    'que requieren un parsing específico.'
)

# 3.3
doc.add_heading('3.3. Investing.com', level=2)
add_paragraph(
    'Del portal financiero Investing.com se descargan series diarias en formato CSV con convención '
    'numérica europea (punto como separador de miles, coma como separador decimal):'
)
add_table_img(
    ['Variable', 'Descripción', 'Obs.'],
    [
        ['bono_es_10y', 'Rentabilidad del Bono español a 10 años', '4.978'],
        ['bono_de_10y', 'Rentabilidad del Bono alemán a 10 años', '4.777'],
        ['spread_es_de', 'Diferencial (spread) España-Alemania 10Y', '2.321'],
        ['vix', 'S&P 500 VIX — índice de volatilidad implícita', '4.989'],
        ['vibex', 'VIBEX — volatilidad implícita del IBEX 35', '1.821'],
        ['vstoxx', 'VSTOXX — vol. implícita del Euro Stoxx 50', '3.117'],
        ['brent', 'Precio del barril de petróleo Brent (USD)', '4.994'],
    ],
    'tab_3_3_investing.png',
    caption='Tabla 2. Variables descargadas de Investing.com.',
    col_widths=[0.15, 0.7, 0.1]
)
add_paragraph(
    'La lectura de estos archivos requiere una función de parsing que elimine símbolos de porcentaje, '
    'reemplace separadores de miles y convierta la coma decimal al punto anglosajón.'
)

# 3.4
doc.add_heading('3.4. Instituto Nacional de Estadística (INE)', level=2)
add_paragraph(
    'Del INE se obtienen dos variables en un formato peculiar de matriz año × mes, donde las filas '
    'representan años y las columnas los meses (M01 a M12):'
)
add_table_img(
    ['Variable', 'Descripción', 'Obs.'],
    [
        ['ipi_yoy', 'Índice de Producción Industrial, var. interanual (%)', '249'],
        ['ipc_sub_mom', 'IPC Subyacente, variación mensual (%)', '250'],
    ],
    'tab_3_4_ine.png',
    caption='Tabla 3. Variables del Instituto Nacional de Estadística.',
    col_widths=[0.15, 0.7, 0.1]
)
add_paragraph(
    'Este formato requiere una transformación de ancho a largo (unpivot/melt), asignando a cada valor '
    'su fecha correspondiente como el primer día del mes (YYYY-MM-01).'
)

# 3.5
doc.add_heading('3.5. Banco Central Europeo (BCE)', level=2)
add_paragraph('Se obtienen dos tipos de datos del BCE:')
p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Euribor 3M y 6M: ')
run_b.bold = True
p.add_run('descargados del ECB Data Portal en formato .xlsx, con columnas DATE y OBS.VALUE. Frecuencia mensual, 250 observaciones cada uno.')
p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Tipos de interés oficiales (tipo_dfr, tipo_mlf, tipo_mro): ')
run_b.bold = True
p.add_run(
    'descargados de epdata.es (portal que replica datos oficiales del BCE) en formato CSV con separador ";" '
    'y decimal ",", con columnas anyo, mes (en nombre español) y los tres tipos. 46 observaciones.'
)

# 3.6
doc.add_heading('3.6. Resumen de fuentes', level=2)
add_paragraph(
    'En total, el proyecto combina 7 fuentes de datos con 4 formatos de archivo distintos '
    '(.xlsx con metadatos Reuters, .xlsx estándar, .csv europeo, .csv con separador ";") y '
    '4 frecuencias temporales (diaria, mensual, trimestral, irregular para tipos BCE).'
)
add_table_img(
    ['Fuente', 'Variables', 'Formato', 'Frecuencia'],
    [
        ['Reuters Eikon (precios)', '35 empresas × 7', 'Excel + metadatos', 'Diaria'],
        ['Reuters Eikon (macro)', '6 variables', 'Excel + metadatos', 'Diaria/Mens./Trim.'],
        ['Investing.com', '7 variables', 'CSV europeo', 'Diaria'],
        ['INE', '2 variables', 'Excel matriz', 'Mensual'],
        ['BCE (ECB Data Portal)', '2 variables', 'Excel estándar', 'Mensual'],
        ['BCE (epdata.es)', '3 variables', 'CSV sep. ";"', 'Irregular'],
    ],
    'tab_3_6_resumen_fuentes.png',
    caption='Tabla 4. Resumen de las 7 fuentes de datos utilizadas.',
    col_widths=[0.28, 0.22, 0.22, 0.22]
)

# ── 4. EXTRACCIÓN Y TRANSFORMACIÓN ────────────────────────────────
doc.add_heading('4. Extracción y transformación de datos', level=1)

doc.add_heading('4.1. ETL de precios de empresas', level=2)
doc.add_heading('4.1.1. Proceso de lectura', level=3)
add_paragraph(
    'Para cada uno de los 35 archivos de precios se implementa la función leer_empresa(filepath, ticker), '
    'que ejecuta los siguientes pasos:'
)
steps = [
    'Lectura sin cabecera (header=None): se lee el archivo completo para detectar la fila donde comienzan los datos reales.',
    'Detección automática de cabecera: se itera fila por fila buscando la cadena "exchange date" (comparación case-insensitive). Esto resuelve el problema de los metadatos de Reuters Eikon, que varían en número de filas entre archivos.',
    'Relectura desde la fila detectada: se relee el archivo con header=fila_cabecera para que pandas interprete correctamente los nombres de columna.',
    'Estandarización de nombres: las columnas originales (Exchange Date, Close, Net, %Chg, Open, Low, High, Volume, Turnover EUR) se renombran al estándar en minúsculas.',
    'Parseo de fechas: la columna fecha se convierte con pd.to_datetime(errors=\'coerce\'), y se eliminan las filas donde la conversión falla.',
    'Filtrado temporal: se descartan las observaciones anteriores a 2005-01-01 para homogeneizar el período de estudio.',
    'Conversión numérica: todas las columnas de precios y volumen se fuerzan a tipo numérico con pd.to_numeric(errors=\'coerce\').',
    'Adición del ticker como columna identificadora.',
    'Ordenación cronológica ascendente por fecha.',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(step)
add_paragraph(
    'Los 35 archivos se procesan en bucle, concatenándose en un único DataFrame consolidado de '
    '157.455 filas × 9 columnas, que cubre el período del 3 de enero de 2005 al 31 de octubre de 2025.'
)

# 4.1.2
doc.add_heading('4.1.2. Heterogeneidad en la cobertura temporal', level=3)
add_paragraph(
    'No todas las empresas cotizan durante los 20 años completos del estudio. Mientras que 19 empresas '
    'disponen de la serie completa (5.323 días de trading desde 2005), otras se incorporaron al IBEX en '
    'fechas posteriores. Los casos más extremos son:'
)
p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Puig Brands (PUIGb.MC): ')
run_b.bold = True
p.add_run('salida a bolsa el 3 de mayo de 2024, con solo 385 observaciones.')
p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Antevenio (ANE.MC): ')
run_b.bold = True
p.add_run('1.112 observaciones desde julio de 2021.')
p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Unicaja (UNI.MC): ')
run_b.bold = True
p.add_run('2.134 observaciones desde junio de 2017.')
add_paragraph(
    'Esta heterogeneidad no se corrige (no se eliminan empresas con menos datos), ya que el formato panel '
    'del dataset maestro permite analizar cada empresa en su rango temporal real.'
)
add_figure('01_cobertura_datos.png', 'Figura 1. Cobertura de datos: porcentaje de nulos por variable (izq.) y heatmap de cobertura anual por variable macro (der.).')

# 4.1.3
doc.add_heading('4.1.3. Variables derivadas', level=3)
add_paragraph('Sobre los datos brutos de precios se calculan dos variables fundamentales para el análisis:')
p = doc.add_paragraph()
run_b = p.add_run('Log-retorno diario (log_ret):')
run_b.bold = True
add_paragraph('    rₜ = ln(Pₜ / Pₜ₋₁)')
add_paragraph(
    'donde Pₜ es el precio de cierre en el día t. El log-retorno se prefiere sobre el retorno aritmético '
    'simple por tres razones: (1) es aditivo en el tiempo, (2) aproxima bien el retorno aritmético para '
    'variaciones pequeñas, y (3) presenta mejores propiedades estadísticas (simetría, estacionariedad). '
    'Se calcula por empresa mediante groupby(\'ticker\').'
)
p = doc.add_paragraph()
run_b = p.add_run('Volatilidad histórica anualizada (vol_hist_21d):')
run_b.bold = True
add_paragraph('    σₜ,₂₁ = √252 · std(rₜ₋₂₀, rₜ₋₁₉, …, rₜ)')
add_paragraph(
    'Corresponde a la desviación típica de los log-retornos en una ventana móvil de 21 días de trading '
    '(aproximadamente un mes), multiplicada por √252 para anualizar la estimación (252 es el número '
    'estándar de días de negociación al año). Se utiliza min_periods=10 para que las empresas con series '
    'cortas no pierdan las primeras 20 observaciones. Esta variable es el objetivo principal (target) de '
    'los modelos de predicción del capítulo de Análisis del Dato.'
)

# 4.2
doc.add_heading('4.2. ETL de variables macroeconómicas', level=2)
doc.add_heading('4.2.1. Funciones de parsing especializadas', level=3)
add_paragraph('La heterogeneidad de formatos requiere cuatro funciones de lectura específicas:')
funcs = [
    ('parsear_numero_europeo(serie)', 'convierte cadenas con formato numérico europeo a float. Elimina el símbolo %, elimina los puntos de separador de miles, reemplaza la coma decimal por punto y aplica pd.to_numeric(errors=\'coerce\'). Se utiliza para todos los archivos de Investing.com.'),
    ('leer_investing_csv(filepath, col_nombre)', 'lee un CSV de Investing.com, detecta automáticamente la columna de cierre (Último/Ultimo), aplica el parsing europeo, convierte la fecha desde formato dd.mm.yyyy y filtra desde 2005.'),
    ('leer_reuters_macro(filepath, col_nombre, sheet)', 'lee un Excel de Reuters Eikon en la hoja especificada (por defecto "First Release Data"). Localiza la fila con el texto "Period", extrae las columnas de período y valor, y parsea los períodos trimestrales ("Q1 2005" → 2005-01-01) y mensuales ("Jan 2005" → 2005-01-01).'),
    ('leer_ine_matriz(filepath, col_nombre)', 'lee un Excel del INE en formato de matriz año × mes. Detecta la fila de cabecera buscando "M01", identifica la columna de años, transforma la matriz a formato largo mediante melt y genera la fecha como el primer día del mes.'),
]
for name, desc in funcs:
    p = doc.add_paragraph()
    run_b = p.add_run(f'{name}: ')
    run_b.bold = True
    p.add_run(desc)

# 4.2.2
doc.add_heading('4.2.2. Organización en bloques temáticos', level=3)
add_paragraph(
    'Las 20 variables macroeconómicas se organizan en cuatro bloques temáticos, cada uno reflejando '
    'un canal de transmisión distinto hacia la volatilidad bursátil:'
)
bloques = [
    ('Bloque 1 — Actividad económica', 'variables que reflejan el estado del ciclo económico español. Incluye el PIB interanual (trimestral), la tasa de paro (trimestral), el Índice de Producción Industrial (mensual) y el PMI (mensual). Capturan el canal real de la economía.'),
    ('Bloque 2 — Condiciones monetarias y financieras', 'variables que reflejan el entorno de tipos de interés y financiación. Incluye bonos soberanos (España y Alemania 10Y), el spread (proxy de riesgo soberano), EUR/USD, Euribor 3M y 6M, tipos oficiales del BCE e IPC interanual. Capturan el canal monetario.'),
    ('Bloque 3 — Precios e inflación', 'el IPC subyacente mensual, que mide las presiones inflacionarias de fondo excluyendo energía y alimentos no elaborados.'),
    ('Bloque 4 — Riesgo global, commodities y divisas', 'variables que capturan el sentimiento de riesgo global. Incluye VIX, VIBEX, VSTOXX, Brent y Gas Natural TTF. Reflejan el canal de contagio entre mercados.'),
]
for name, desc in bloques:
    p = doc.add_paragraph()
    run_b = p.add_run(f'{name}: ')
    run_b.bold = True
    p.add_run(desc)

# 4.2.3
doc.add_heading('4.2.3. Carga en base de datos', level=3)
add_paragraph('Las variables se agrupan en 5 tablas SQLite según su frecuencia y temática:')
add_table_img(
    ['Tabla', 'Frecuencia', 'Variables', 'Filas'],
    [
        ['macro_act_mensual', 'Mensual', 'ipi_yoy, pmi, ipc_sub_mom', '250'],
        ['macro_act_trimes', 'Trimestral', 'pib_yoy, tasa_paro', '82'],
        ['macro_mon_diario', 'Diaria', 'bono_es_10y, bono_de_10y, spread_es_de, eur_usd', '5.865'],
        ['macro_mon_mensual', 'Mensual', 'euribor_3m, euribor_6m, tipo_dfr, tipo_mlf, tipo_mro, ipc_yoy', '423'],
        ['macro_riesgo', 'Diaria', 'vix, vibex, vstoxx, brent, gas_ttf', '5.091'],
    ],
    'tab_4_2_tablas_sqlite.png',
    caption='Tabla 5. Tablas macroeconómicas en SQLite.',
    col_widths=[0.22, 0.13, 0.52, 0.08]
)
add_paragraph(
    'Las tablas dentro de cada bloque se construyen mediante merge(on=\'fecha\', how=\'outer\') para '
    'preservar todas las fechas disponibles, aunque provengan de series con rangos temporales distintos.'
)

# ── 5. LIMPIEZA ────────────────────────────────────────────────────
doc.add_heading('5. Limpieza de datos', level=1)

doc.add_heading('5.1. Datos de precios — Detección y tratamiento de nulos', level=2)
add_paragraph(
    'El análisis de calidad del dataset de precios revela 8 valores nulos en la columna close, '
    'distribuidos en 4 empresas:'
)
add_table_img(
    ['Empresa', 'Ticker', 'Nulos en close'],
    [
        ['CaixaBank', 'CABK.MC', '1'],
        ['Ferrovial', 'FER.MC', '1'],
        ['Merlin Properties', 'MRL.MC', '1'],
        ['Laboratorios Rovi', 'ROVI.MC', '5'],
    ],
    'tab_5_1_nulos_close.png',
    caption='Tabla 6. Empresas con valores nulos en el precio de cierre.',
    col_widths=[0.35, 0.3, 0.25]
)
add_paragraph(
    'Estos nulos corresponden a días en los que la plataforma Reuters no registró precio de cierre, '
    'probablemente por suspensiones temporales de cotización o errores en la exportación de datos.'
)
p = doc.add_paragraph()
run_b = p.add_run('Decisión: interpolación lineal. ')
run_b.bold = True
p.add_run(
    'Se opta por pandas.interpolate(method=\'linear\', limit_direction=\'both\'), aplicada por empresa. '
    'Este método estima el precio faltante como el punto medio ponderado entre el último precio conocido '
    'y el siguiente disponible. Se prefiere la interpolación lineal sobre las alternativas por las siguientes razones:'
)
reasons = [
    'Frente a forward-fill (ffill): los precios tienen tendencia, por lo que arrastrar el último valor conocido introduce un sesgo sistemático.',
    'Frente a eliminación de filas: eliminar el día completo penalizaría a las otras 34 empresas que sí tienen dato ese día.',
    'Frente a la media: el precio medio histórico no tiene sentido para una serie con tendencia temporal.',
]
for r in reasons:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(r)
add_paragraph(
    'Tras la imputación, se recalculan log_ret y vol_hist_21d para que reflejen los valores interpolados.'
)
add_figure('03_interpolacion_nulos.png', 'Figura 3. Resultado de la interpolación lineal en las 4 empresas con nulos en el precio de cierre.')
add_paragraph(
    'Verificación final: 0 nulos en close tras la imputación. 0 duplicados por par (ticker, fecha). 0 precios negativos o cero.',
    bold=True
)

# 5.2
doc.add_heading('5.2. Datos macroeconómicos — Nulos estructurales', level=2)
add_paragraph(
    'En las variables macroeconómicas, los nulos no son errores sino ausencias estructurales derivadas de:'
)
p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Distintos rangos temporales: ')
run_b.bold = True
p.add_run('el VIBEX comienza en 2015, el spread España-Alemania en 2014, los tipos BCE en 2009.')
p = doc.add_paragraph(style='List Bullet')
run_b = p.add_run('Distinta frecuencia de publicación: ')
run_b.bold = True
p.add_run('el PIB solo se publica una vez al trimestre; entre publicaciones, los valores permanecen constantes.')
add_paragraph(
    'Estas ausencias se resuelven en el proceso de alineación temporal (sección 6) mediante forward-fill, '
    'que es la aproximación correcta desde el punto de vista de la información disponible para un inversor.'
)

# ── 6. ALINEACIÓN TEMPORAL ─────────────────────────────────────────
doc.add_heading('6. Alineación temporal y construcción del dataset maestro', level=1)

doc.add_heading('6.1. El problema de las frecuencias múltiples', level=2)
add_paragraph(
    'El principal reto técnico de este trabajo es la integración de datos con frecuencias muy distintas '
    'en un único dataset de frecuencia diaria:'
)
freqs = [
    'Datos diarios: precios de empresas (5.323 fechas), bonos, VIX, Brent, EUR/USD.',
    'Datos mensuales: IPI, IPC, Euribor, tipos BCE (12 observaciones/año).',
    'Datos trimestrales: PIB, tasa de paro (4 observaciones/año).',
    'Datos irregulares: tipos de interés del BCE (solo cuando hay decisión de política monetaria).',
]
for f in freqs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f)
add_paragraph(
    'No es posible simplemente hacer un merge por fecha, porque las fechas de las variables mensuales o '
    'trimestrales rara vez coinciden con días de negociación bursátil.'
)

# 6.2
doc.add_heading('6.2. Solución: forward-fill y merge_asof', level=2)
add_paragraph('La estrategia adoptada sigue el enfoque estándar de la econometría financiera:')
merge_steps = [
    'Construcción del calendario bursátil: se extraen las 5.323 fechas de negociación únicas de la tabla precios_empresas.',
    'Forward-fill interno (ffill()) en cada tabla macro: propaga el último valor publicado hasta la siguiente publicación.',
    'Alineación con pd.merge_asof: para cada fecha del calendario bursátil, se asigna el último valor macro disponible (direction=\'backward\').',
    'Combinación de tablas macro: las 5 tablas alineadas se fusionan en una tabla df_macro_diario de 5.323 filas × 21 columnas.',
    'Merge final con precios: cada empresa en cada día recibe sus variables macro correspondientes.',
]
for i, step in enumerate(merge_steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(step)
add_paragraph(
    'Resultado: el dataset maestro contiene 157.455 filas × 31 columnas (35 empresas × ~4.500 días).',
    bold=True
)

# 6.3
doc.add_heading('6.3. Estructura del dataset maestro', level=2)
add_table_img(
    ['Grupo', 'Variables', 'Tipo'],
    [
        ['Identificación', 'fecha, ticker', 'Clave compuesta'],
        ['Precios', 'close, open, high, low, volume', 'Numérica continua'],
        ['Retornos', 'net, pct_chg, log_ret', 'Numérica continua'],
        ['Volatilidad', 'vol_hist_21d', 'Numérica (target)'],
        ['Monetario diario', 'bono_es_10y, bono_de_10y, spread_es_de, eur_usd', 'Numérica continua'],
        ['Riesgo global', 'vix, vibex, vstoxx, brent, gas_ttf', 'Numérica continua'],
        ['Actividad mensual', 'ipi_yoy, pmi, ipc_sub_mom', 'Numérica continua'],
        ['Monetario mensual', 'euribor_3m, euribor_6m, tipo_dfr, tipo_mlf, tipo_mro, ipc_yoy', 'Numérica continua'],
        ['Actividad trimestral', 'pib_yoy, tasa_paro', 'Numérica continua'],
    ],
    'tab_6_3_estructura_maestro.png',
    caption='Tabla 7. Estructura del dataset maestro (157.455 filas × 31 columnas).',
    col_widths=[0.2, 0.55, 0.2]
)

# 6.4
doc.add_heading('6.4. Cobertura del dataset maestro', level=2)
add_paragraph('El análisis de nulos del dataset maestro revela la cobertura efectiva de cada variable:')
add_table_img(
    ['Variable', 'Cobertura', 'Período efectivo'],
    [
        ['euribor_3m', '99,8%', '2005 – 2025'],
        ['pib_yoy', '99,2%', '2005 – 2025'],
        ['vix', '96,2%', '2005 – 2025'],
        ['brent', '95,1%', '2005 – 2025'],
        ['bono_es_10y', '89,7%', '2005 – 2025'],
        ['ipc_yoy', '70,9%', '2008 – 2025'],
        ['spread_es_de', '52,6%', '2014 – 2025'],
        ['vibex', '39,1%', '2015 – 2025'],
        ['pmi', '16,8%', '2023 – 2025'],
    ],
    'tab_6_4_cobertura.png',
    caption='Tabla 8. Cobertura de las variables macroeconómicas en el dataset maestro.',
    highlight_col=1,
    highlight_vals={
        '99,8%': '#D4EDDA', '99,2%': '#D4EDDA', '96,2%': '#D4EDDA',
        '95,1%': '#D4EDDA', '89,7%': '#D4EDDA',
        '70,9%': '#FFF3CD',
        '52,6%': '#FDDCDC', '39,1%': '#FDDCDC', '16,8%': '#FDDCDC',
    },
    col_widths=[0.25, 0.2, 0.35]
)

# ── 7. EDA ─────────────────────────────────────────────────────────
doc.add_heading('7. Análisis exploratorio de datos (EDA)', level=1)
add_paragraph(
    'El análisis exploratorio se estructura en secciones que progresan desde la descripción básica '
    'hasta los tests formales de hipótesis, documentando cada hallazgo y su implicación para la modelización.'
)

# 7.1
doc.add_heading('7.1. Estadísticos descriptivos', level=2)
add_paragraph('Los log-retornos del panel completo (157.420 observaciones) presentan las siguientes características:')
add_table_img(
    ['Estadístico', 'Valor'],
    [
        ['Media', '0,00015'],
        ['Desviación típica', '0,0212'],
        ['Asimetría (skewness)', '-0,071'],
        ['Curtosis (exceso)', '16,00'],
        ['Mínimo', '-0,4297'],
    ],
    'tab_7_1_descriptivos.png',
    caption='Tabla 9. Estadísticos descriptivos de los log-retornos.',
    col_widths=[0.45, 0.35]
)
add_paragraph(
    'La curtosis de 16 (frente al valor de 3 de una distribución normal) indica colas extremadamente '
    'pesadas: los eventos extremos son mucho más frecuentes de lo que predeciría una distribución '
    'gaussiana. Este hallazgo, conocido como fat tails, implica que los modelos que asuman normalidad '
    'subestimarán sistemáticamente el riesgo de cola.'
)
add_paragraph(
    'La volatilidad media del IBEX 35 es de 0,2926 (29,26% anualizada), con una distribución marcadamente '
    'asimétrica a la derecha (asimetría = 2,44, curtosis = 12,53).'
)

# 7.2
doc.add_heading('7.2. Transformación a log-retornos: antes y después', level=2)
add_paragraph(
    'La siguiente figura muestra la comparación antes/después para Santander (SAN.MC): la serie de precios '
    'exhibe tendencias y no estacionariedad, mientras que los log-retornos oscilan alrededor de cero con '
    'varianza variable en el tiempo (volatility clustering).'
)
add_figure('02_antes_despues_logretos.png', 'Figura 2. Antes y después: precios vs. log-retornos para Santander (SAN.MC).')

# 7.3
doc.add_heading('7.3. Distribución de los log-retornos', level=2)
add_paragraph(
    'Histograma conjunto de los 157.420 log-retornos con curva normal teórica y QQ-plot. '
    'La desviación en ambas colas del QQ-plot confirma la no normalidad.'
)
add_figure('04_distribucion_logreturns.png', 'Figura 4. Distribución de log-retornos: histograma (izq.) y QQ-plot (der.).')

# 7.4
doc.add_heading('7.4. Evolución temporal de la volatilidad', level=2)
add_paragraph(
    'Serie temporal de la volatilidad media del IBEX 35 con los cuatro episodios de crisis marcados:'
)
crises = [
    'Lehman Brothers (septiembre 2008): volatilidad del 80-100% anualizado.',
    'Crisis de deuda europea (mayo 2010): prima de riesgo española disparándose.',
    'COVID-19 (marzo 2020): pico más abrupto, volatilidad ×4 en dos semanas.',
    'Guerra de Ucrania (febrero 2022): pico moderado pero sostenido.',
]
for c in crises:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(c)
add_paragraph(
    'Este gráfico evidencia el fenómeno de volatility clustering (Mandelbrot, 1963): los períodos de alta '
    'volatilidad se agrupan y persisten, justificando modelos autorregresivos como el GARCH.'
)
add_figure('05_volatilidad_temporal.png', 'Figura 5. Evolución temporal de la volatilidad media del IBEX 35 (2005-2025).')

# 7.5
doc.add_heading('7.5. Volatilidad por empresa y sector', level=2)
add_paragraph(
    'Boxplot por empresa y barras por sector GICS. Real Estate (35,4%) y Materials (34,2%) son los '
    'sectores más volátiles; Communication Services (23,5%) el más estable.'
)
add_figure('06_volatilidad_empresa_sector.png', 'Figura 6. Volatilidad por empresa (izq.) y por sector GICS (der.).')

# 7.6
doc.add_heading('7.6. Evolución de las variables macroeconómicas', level=2)
add_paragraph(
    'Cuadrícula 4×2 con las 8 variables macro más relevantes (2005-2025). Permite identificar regímenes: '
    'era de tipos bajos (2012-2021), crisis de inflación (2021-2023), normalización monetaria del BCE.'
)
add_figure('07_macro_evolucion.png', 'Figura 7. Evolución temporal de las principales variables macroeconómicas.')

# 7.7
doc.add_heading('7.7. Tests de estacionariedad (ADF)', level=2)
add_paragraph(
    'La estacionariedad es un requisito fundamental para los modelos econométricos. Test ADF con '
    'selección automática de retardos por criterio AIC:'
)
add_table_img(
    ['Variable', 'ADF stat', 'p-valor', 'Estacionaria'],
    [
        ['Precio cierre (SAN.MC)', '-1,73', '0,4166', 'NO'],
        ['Log-retorno (SAN.MC)', '-43,24', '0,0000', 'SÍ'],
        ['Volatilidad 21d (SAN.MC)', '-6,17', '0,0000', 'SÍ'],
        ['VIX', '-5,60', '0,0000', 'SÍ'],
        ['Bono ES 10Y', '-1,36', '0,6005', 'NO'],
        ['Euribor 3M', '-1,91', '0,3255', 'NO'],
        ['IPC YoY', '-1,72', '0,4230', 'NO'],
        ['PIB YoY', '-4,48', '0,0002', 'SÍ'],
    ],
    'tab_7_7_adf.png',
    caption='Tabla 10. Resultados del test Augmented Dickey-Fuller.',
    highlight_col=3,
    highlight_vals={'SÍ': '#D4EDDA', 'NO': '#FDDCDC'},
    col_widths=[0.35, 0.15, 0.15, 0.18]
)
add_paragraph('Implicaciones para la modelización:', bold=True)
impls = [
    'Los precios no se utilizan como features, confirmando la necesidad de la transformación a log-retornos.',
    'La volatilidad y los log-retornos son estacionarios, validando su uso como variable objetivo.',
    'Las variables macro no estacionarias (bono, euribor, IPC) se usan con precaución, coherente con su papel de "información disponible".',
]
for im in impls:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(im)

# 7.8
doc.add_heading('7.8. Autocorrelación y memoria larga', level=2)
add_paragraph(
    'Las funciones ACF y PACF con 60 retardos revelan persistencia extraordinaria: la ACF es significativa '
    'en todos los lags hasta 60, confirmando la memoria larga. La PACF muestra picos en lags 1, 5 y 21 '
    '(escalas diaria, semanal y mensual).'
)
add_paragraph(
    'Nota metodológica: la ventana de 21 días introduce autocorrelación mecánica en los primeros 20 lags. '
    'La persistencia genuina se observa en lags > 21.',
    italic=True
)
add_paragraph(
    'Esta memoria larga justifica el modelo HAR (Corsi, 2009) con retardos a tres escalas (1, 5, 21 días) '
    'y el modelo GARCH para la persistencia autorregresiva de la varianza condicional.'
)
add_figure('08_acf_pacf_volatilidad.png', 'Figura 8. ACF (arriba) y PACF (abajo) de la volatilidad media del IBEX 35, 60 retardos.')

# 7.9
doc.add_heading('7.9. Correlaciones con la volatilidad', level=2)
add_paragraph('Correlaciones de Pearson entre cada variable macro y la volatilidad media del IBEX 35:')
add_paragraph('Correlaciones positivas más fuertes:', bold=True)
pos_corrs = [
    'VIBEX: +0,824 (volatilidad implícita del propio IBEX)',
    'VIX: +0,754 (canal de contagio global)',
    'VSTOXX: +0,728 (canal europeo)',
    'Bono ES 10Y: +0,208 (riesgo soberano)',
    'Tasa de paro: +0,178 (debilidad económica)',
]
for c in pos_corrs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(c)
add_paragraph('Correlaciones negativas más fuertes:', bold=True)
neg_corrs = [
    'IPI YoY: -0,442 (mayor actividad industrial → menor incertidumbre)',
    'PIB YoY: -0,429 (crecimiento económico → menor volatilidad)',
]
for c in neg_corrs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(c)
add_paragraph(
    'El heatmap triangular revela multicolinealidades importantes (VIX-VSTOXX: +0,93; euribor_3m-euribor_6m: +0,99).'
)
add_figure('09_correlaciones_macro_vol.png', 'Figura 9. Correlaciones macro-volatilidad (izq.) y heatmap entre variables macro (der.).')

# 7.10
doc.add_heading('7.10. Relación dinámica: volatilidad, VIX y Euribor', level=2)
add_paragraph(
    'Paneles temporales con ejes duales: (arriba) volatilidad IBEX y VIX con sincronización casi perfecta '
    'en crisis; (abajo) volatilidad IBEX y Euribor 3M, relación que refleja las decisiones del BCE.'
)
add_figure('10_volatilidad_vix_euribor.png', 'Figura 10. Volatilidad IBEX 35 vs. VIX (arriba) y vs. Euribor 3M (abajo).')

# 7.11
doc.add_heading('7.11. Cross-correlación con retardos', level=2)
add_paragraph('Cross-correlogramas de 6 variables macro clave con la volatilidad IBEX (lags 0 a 60):')
add_table_img(
    ['Variable', 'CCF máxima', 'Lag óptimo', 'Interpretación'],
    [
        ['VIX', '+0,754', '0', 'Contemporánea'],
        ['Bono ES 10Y', '+0,208', '0', 'Contemporánea'],
        ['Spread ES-DE', '-0,119', '59', 'Adelanta la volatilidad'],
        ['Euribor 3M', '-0,048', '59', 'Efecto retardado'],
        ['Brent', '-0,056', '19', 'Efecto retardado'],
        ['IPC YoY', '-0,036', '14', 'Efecto retardado'],
    ],
    'tab_7_11_ccf.png',
    caption='Tabla 11. Cross-correlaciones máximas entre variables macro y volatilidad IBEX.',
    col_widths=[0.22, 0.18, 0.15, 0.35]
)
add_paragraph(
    'El VIX y el bono español tienen máxima correlación en lag 0, mientras que spread, euribor y brent '
    'la presentan en lags positivos, sugiriendo poder predictivo anticipado.'
)
add_figure('13_ccf_macro_volatilidad.png', 'Figura 13. Cross-correlogramas de variables macro vs. volatilidad IBEX (lags 0-60).')

# 7.12
doc.add_heading('7.12. Detección de outliers', level=2)
add_paragraph('Se aplican dos métodos complementarios:')
p = doc.add_paragraph()
run_b = p.add_run('Método IQR: ')
run_b.bold = True
p.add_run('outlier si x < Q₁ - 1,5·IQR o x > Q₃ + 1,5·IQR.')
p = doc.add_paragraph()
run_b = p.add_run('Método Z-score: ')
run_b.bold = True
p.add_run('outlier si |z| > 3.')
add_table_img(
    ['Variable', 'Outliers IQR', '% IQR', 'Outliers Z>3', '% Z>3'],
    [
        ['vol_hist_21d', '8.169', '5,20%', '2.801', '1,78%'],
        ['log_ret', '9.500', '6,03%', '2.446', '1,55%'],
        ['vix', '6.774', '4,47%', '3.157', '2,08%'],
    ],
    'tab_7_12_outliers.png',
    caption='Tabla 12. Detección de outliers por IQR y Z-score.',
    col_widths=[0.22, 0.18, 0.12, 0.18, 0.12]
)
p = doc.add_paragraph()
run_b = p.add_run('Decisión: mantener todos los outliers. ')
run_b.bold = True
p.add_run(
    'Los valores extremos corresponden a crisis reales (2008, 2010, 2020, 2022) y son señal económica, '
    'no errores de medición. El GARCH con t-Student es robusto a valores extremos por diseño.'
)
add_figure('11_outliers_iqr_zscore.png', 'Figura 11. Outliers detectados por IQR y Z-score en volatilidad, log-retornos y VIX.')

# 7.13
doc.add_heading('7.13. Test de normalidad Jarque-Bera', level=2)
add_paragraph('QQ-plots de las series principales con resultados del test Jarque-Bera:')
add_table_img(
    ['Variable', 'JB', 'p-valor', 'Normal', 'Curtosis'],
    [
        ['vol_hist_21d', '25.346', '<0,0001', 'NO', '12,53'],
        ['log_ret', '1.678.375', '<0,0001', 'NO', '19,00'],
        ['vix', '23.312', '<0,0001', 'NO', '12,31'],
        ['bono_es_10y', '199', '<0,0001', 'NO', '2,04'],
        ['euribor_3m', '638', '<0,0001', 'NO', '2,25'],
        ['brent', '158', '<0,0001', 'NO', '2,29'],
    ],
    'tab_7_13_jarque_bera.png',
    caption='Tabla 13. Resultados del test de normalidad Jarque-Bera.',
    highlight_col=3,
    highlight_vals={'NO': '#FDDCDC'},
    col_widths=[0.22, 0.2, 0.15, 0.12, 0.15]
)
add_paragraph('Ninguna variable financiera sigue una distribución normal. Esto justifica:', bold=True)
jb_impls = [
    'Utilizar la distribución t-Student para las innovaciones del modelo GARCH.',
    'Interpretar los intervalos de confianza del OLS con cautela, aunque la estimación puntual por MCO sigue siendo consistente (Gauss-Markov).',
]
for j in jb_impls:
    p = doc.add_paragraph(style='List Number')
    p.add_run(j)
add_figure('12_jarque_bera_qqplots.png', 'Figura 12. QQ-plots contra distribución normal con resultados del test Jarque-Bera.')

# 7.14
doc.add_heading('7.14. Test de Kruskal-Wallis entre sectores', level=2)
add_paragraph(
    'El test de Kruskal-Wallis evalúa si la distribución de volatilidad difiere entre los 11 sectores GICS:'
)
add_paragraph('    H = 7.259,89    (p ≈ 0,000000)', bold=True)
add_paragraph(
    'Se rechaza la igualdad de distribuciones. El post-hoc con Mann-Whitney U y corrección de Bonferroni '
    '(α = 0,0009 para 55 comparaciones) identifica 47 de 55 pares significativamente distintos. '
    'Los sectores más volátiles: Real Estate (35,4%), Materials (34,2%), Financials (30,7%). '
    'Los menos volátiles: Communication Services (23,5%), Consumer Discretionary (24,9%).'
)
add_figure('14_kruskal_wallis_sectores.png', 'Figura 14. Boxplot de volatilidad por sector GICS con resultado del test Kruskal-Wallis.')

# 7.15
doc.add_heading('7.15. Split temporal train/test', level=2)
add_paragraph(
    'Se adopta un split fijo por fecha (no aleatorio) para respetar la estructura temporal. Las observaciones '
    'más recientes conforman el conjunto de test, evitando data leakage.'
)
add_figure('15_split_temporal.png', 'Figura 15. Visualización del split temporal train/test.')

# ── 8. DECISIONES ──────────────────────────────────────────────────
doc.add_heading('8. Decisiones de diseño para la modelización', level=1)
add_paragraph('El EDA culmina con una tabla de decisión que consolida los hallazgos:')
add_table_img(
    ['Decisión', 'Variables', 'Justificación'],
    [
        ['MANTENER (20)', 'Todas excepto las siguientes', 'Cobertura >70%, relevancia económica, correlación significativa'],
        ['REDUNDANTE (2)', 'tipo_mlf, tipo_mro', 'Colinealidad perfecta con tipo_dfr (BCE los anuncia simultáneamente)'],
        ['USO LIMITADO (1)', 'pmi', 'Solo 3 años de datos (2023-2025), 83,2% de nulos'],
    ],
    'tab_8_decisiones.png',
    caption='Tabla 14. Tabla de decisión de variables para la fase de modelización.',
    highlight_col=0,
    highlight_vals={
        'MANTENER (20)': '#D4EDDA',
        'REDUNDANTE (2)': '#FFF3CD',
        'USO LIMITADO (1)': '#FDDCDC',
    },
    col_widths=[0.2, 0.3, 0.45]
)

# ── 9. BASE DE DATOS FINAL ────────────────────────────────────────
doc.add_heading('9. Base de datos final', level=1)
add_paragraph('La base de datos SQLite (tfg.db) resultante tiene la siguiente estructura:')
add_table_img(
    ['Tabla', 'Filas', 'Cols.', 'Descripción'],
    [
        ['precios_empresas', '157.455', '11', 'Precios OHLCV + retornos + volatilidad'],
        ['ref_empresas', '36', '5', 'Catálogo de empresas IBEX 35'],
        ['macro_act_mensual', '250', '4', 'Actividad económica mensual'],
        ['macro_act_trimes', '82', '3', 'Actividad económica trimestral'],
        ['macro_mon_diario', '5.865', '5', 'Condiciones monetarias diarias'],
        ['macro_mon_mensual', '423', '7', 'Condiciones monetarias mensuales'],
        ['macro_riesgo', '5.091', '6', 'Riesgo global y commodities'],
        ['dataset_maestro', '157.455', '31', 'Dataset consolidado para modelización'],
    ],
    'tab_9_bd_final.png',
    caption='Tabla 15. Estructura de la base de datos tfg.db (8 tablas, >325.000 registros).',
    col_widths=[0.22, 0.1, 0.07, 0.5]
)
add_paragraph(
    'El diseño en estrella permite trazabilidad completa del dato y eficiencia en la modelización. '
    'Se generan copias de seguridad en CSV y Parquet en data/processed/.'
)

# ── 10. RESUMEN FIGURAS ───────────────────────────────────────────
doc.add_heading('10. Resumen de figuras generadas', level=1)
add_table_img(
    ['Fig.', 'Título', 'Sección'],
    [
        ['1', 'Cobertura de datos y porcentaje de nulos', '4.1.2'],
        ['2', 'Antes y después: precios vs. log-retornos', '7.2'],
        ['3', 'Resultado de la interpolación lineal de nulos', '5.1'],
        ['4', 'Distribución de log-retornos y QQ-plot', '7.3'],
        ['5', 'Evolución temporal de la volatilidad del IBEX 35', '7.4'],
        ['6', 'Volatilidad por empresa y sector GICS', '7.5'],
        ['7', 'Evolución de variables macroeconómicas', '7.6'],
        ['8', 'ACF y PACF de la volatilidad', '7.8'],
        ['9', 'Correlaciones macro-volatilidad', '7.9'],
        ['10', 'Volatilidad IBEX vs. VIX y Euribor', '7.10'],
        ['11', 'Detección de outliers (IQR y Z-score)', '7.12'],
        ['12', 'QQ-plots y test Jarque-Bera', '7.13'],
        ['13', 'Cross-correlogramas macro-volatilidad', '7.11'],
        ['14', 'Test Kruskal-Wallis por sector GICS', '7.14'],
        ['15', 'Split temporal train/test', '7.15'],
    ],
    'tab_10_resumen_figuras.png',
    caption='Tabla 16. Resumen de las 15 figuras generadas en la fase de Ingeniería del Dato.',
    col_widths=[0.06, 0.68, 0.1]
)

# ── GUARDAR ────────────────────────────────────────────────────────
doc.save(OUT)
print(f"✓ Documento guardado en:\n  {OUT}")
print(f"✓ Tablas renderizadas en:\n  {TABS}")
print(f"✓ Total tablas imagen: {len(os.listdir(TABS))}")
